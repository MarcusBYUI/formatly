"""
FormattingAnalyzer Module for Formatly V3
Performs in-depth style compliance checks for academic documents.
Analyzes citations, headings, references, margins, spacing, and other formatting elements.
"""

import re
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from style_guides import STYLE_GUIDES

class FormattingAnalyzer:
    """
    Analyzes document formatting for compliance with academic style guides.
    Performs checks on citations, headings, references, margins, spacing, etc.
    """
    
    def __init__(self, style_name: str = "apa"):
        """
        Initialize the formatting analyzer.
        
        Args:
            style_name: Name of the style guide to use (default: "apa")
        """
        self.style_name = style_name.lower()
        self.style_guide = STYLE_GUIDES.get(self.style_name, STYLE_GUIDES["apa"])
        self.issue_categories = [
            "citations", "headings", "references", "margins", 
            "spacing", "font", "paragraphs", "title_page"
        ]
    
    def analyze_document(self, doc_path: str) -> Dict:
        """
        Perform comprehensive analysis of document formatting.
        
        Args:
            doc_path: Path to the Word document file
            
        Returns:
            Dictionary containing analysis results
        """
        try:
            doc = Document(doc_path)
            
            # Compile all analysis results
            return {
                "document_info": self._get_document_info(doc),
                "style_requirements": self._get_style_requirements(),
                "compliance_score": self._calculate_compliance_score(doc),
                "formatting_issues": self._find_formatting_issues(doc),
                "citation_analysis": self._analyze_citations(doc),
                "reference_analysis": self._analyze_references(doc),
                "heading_structure": self._analyze_heading_structure(doc),
                "spacing_analysis": self._analyze_spacing(doc),
                "margin_analysis": self._analyze_margins(doc)
            }
        except Exception as e:
            return {
                "error": f"Analysis failed: {str(e)}",
                "document_info": {"path": doc_path, "style": self.style_name},
                "compliance_score": {"overall": 0, "details": {}},
                "formatting_issues": []
            }
    
    def _get_document_info(self, doc: Document) -> Dict:
        """Extract basic document information."""
        sections = len(doc.sections)
        paragraphs = len(doc.paragraphs)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
        
        return {
            "style": self.style_name.upper(),
            "sections": sections,
            "paragraphs": paragraphs,
            "word_count": word_count,
            "has_title_page": self._has_title_page(doc)
        }
    
    def _get_style_requirements(self) -> Dict:
        """Extract key style requirements for reference."""
        meta = self.style_guide.get("meta", {})
        margins = self.style_guide.get("margins", {})
        
        return {
            "font": meta.get("default_font", "Times New Roman"),
            "requires_title_page": meta.get("title_page", False),
            "requires_abstract": meta.get("requires_abstract", False),
            "line_spacing": "Double",  # Default for most academic styles
            "margins": {
                "left": self._convert_inches_to_str(margins.get("left", Inches(1))),
                "right": self._convert_inches_to_str(margins.get("right", Inches(1))),
                "top": self._convert_inches_to_str(margins.get("top", Inches(1))),
                "bottom": self._convert_inches_to_str(margins.get("bottom", Inches(1)))
            },
            "citation_format": meta.get("citation_format", "({Author}, {Year})")
        }
    
    def _convert_inches_to_str(self, inches_obj) -> str:
        """Convert docx Inches object to string representation."""
        try:
            if hasattr(inches_obj, "inches"):
                return f"{inches_obj.inches:.1f}\""
            else:
                return f"{inches_obj:.1f}\""
        except:
            return "1.0\""  # Default fallback
    
    def _has_title_page(self, doc: Document) -> bool:
        """Check if document appears to have a title page."""
        if not doc.paragraphs:
            return False
        
        # Check if first page contains typical title page elements
        title_indicators = ["title:", "author:", "course:", "professor:", "date:"]
        first_page_text = " ".join([p.text.lower() for p in doc.paragraphs[:10]])
        
        # Count how many title indicators are present
        indicator_count = sum(1 for indicator in title_indicators if indicator in first_page_text)
        
        # If at least 2 indicators are present, it's likely a title page
        return indicator_count >= 2
    
    def _calculate_compliance_score(self, doc: Document) -> Dict:
        """Calculate overall formatting compliance score."""
        # Initialize scores for different categories
        category_scores = {
            "citations": 0,
            "headings": 0,
            "references": 0,
            "margins": 0,
            "spacing": 0,
            "font": 0,
            "paragraphs": 0,
            "title_page": 0
        }
        
        # Analyze each category and assign scores (0-100)
        category_scores["margins"] = self._score_margins(doc)
        category_scores["spacing"] = self._score_spacing(doc)
        category_scores["headings"] = self._score_headings(doc)
        category_scores["font"] = self._score_fonts(doc)
        category_scores["paragraphs"] = self._score_paragraphs(doc)
        category_scores["references"] = self._score_references(doc)
        category_scores["citations"] = self._score_citations(doc)
        category_scores["title_page"] = self._score_title_page(doc)
        
        # Calculate weighted overall score
        weights = {
            "citations": 0.15,
            "headings": 0.15,
            "references": 0.15,
            "margins": 0.10,
            "spacing": 0.10,
            "font": 0.15,
            "paragraphs": 0.10,
            "title_page": 0.10
        }
        
        overall_score = sum(score * weights[category] for category, score in category_scores.items())
        
        # Assign grade based on score
        grade = self._score_to_grade(overall_score)
        
        return {
            "overall": round(overall_score, 1),
            "grade": grade,
            "details": category_scores
        }
    
    def _score_to_grade(self, score: float) -> str:
        """Convert numerical score to letter grade."""
        if score >= 97:
            return "A+"
        elif score >= 93:
            return "A"
        elif score >= 90:
            return "A-"
        elif score >= 87:
            return "B+"
        elif score >= 83:
            return "B"
        elif score >= 80:
            return "B-"
        elif score >= 77:
            return "C+"
        elif score >= 73:
            return "C"
        elif score >= 70:
            return "C-"
        elif score >= 67:
            return "D+"
        elif score >= 63:
            return "D"
        elif score >= 60:
            return "D-"
        else:
            return "F"
    
    def _score_margins(self, doc: Document) -> int:
        """Score margin compliance."""
        try:
            if not doc.sections:
                return 50  # Default middle score if no sections
            
            # Get required margins from style guide
            required_margins = self.style_guide.get("margins", {})
            
            # Get actual margins from first section
            section = doc.sections[0]
            actual_margins = {
                "left": section.left_margin,
                "right": section.right_margin,
                "top": section.top_margin,
                "bottom": section.bottom_margin
            }
            
            # Calculate margin compliance (allow small deviations)
            compliant_margins = 0
            for margin_type, required in required_margins.items():
                if margin_type in actual_margins:
                    # Allow 0.1 inch deviation
                    if abs(actual_margins[margin_type].inches - required.inches) <= 0.1:
                        compliant_margins += 1
            
            # Calculate score (percentage of compliant margins)
            margin_types = ["left", "right", "top", "bottom"]
            score = (compliant_margins / len(margin_types)) * 100
            
            return int(score)
        except:
            return 50  # Default middle score if error
    
    def _score_spacing(self, doc: Document) -> int:
        """Score line spacing compliance."""
        try:
            # Most academic styles require double spacing
            required_spacing = WD_LINE_SPACING.DOUBLE
            
            # Check paragraphs for proper spacing
            if not doc.paragraphs:
                return 50  # Default middle score if no paragraphs
            
            # Count paragraphs with correct spacing
            correct_spacing = 0
            total_checked = 0
            
            for p in doc.paragraphs:
                if p.text.strip():  # Only check non-empty paragraphs
                    total_checked += 1
                    try:
                        if p.paragraph_format.line_spacing_rule == required_spacing:
                            correct_spacing += 1
                    except:
                        pass  # Skip if cannot determine spacing
            
            if total_checked == 0:
                return 50  # Default middle score if no valid paragraphs
            
            # Calculate score
            score = (correct_spacing / total_checked) * 100
            return int(score)
        except:
            return 50  # Default middle score if error
    
    def _score_headings(self, doc: Document) -> int:
        """Score heading structure compliance."""
        try:
            # Check if document has proper heading structure
            headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
            
            if not headings:
                return 30  # Penalize for no headings
            
            # Check for logical heading structure (no skipping levels)
            heading_levels = [int(h.style.name.split(' ')[1]) for h in headings 
                             if len(h.style.name.split(' ')) > 1 and h.style.name.split(' ')[1].isdigit()]
            
            if not heading_levels:
                return 40  # Some heading-like structure, but not standard headings
            
            # Check if heading levels are used in proper sequence
            level_issues = 0
            for i in range(len(heading_levels) - 1):
                # Heading should not jump more than one level
                if heading_levels[i+1] - heading_levels[i] > 1:
                    level_issues += 1
            
            # Calculate percentage of well-structured headings
            if len(heading_levels) <= 1:
                structure_score = 70  # Single heading is okay but not ideal
            else:
                structure_score = 100 - (level_issues / (len(heading_levels) - 1)) * 60
            
            # Check heading formatting
            format_score = self._check_heading_formatting(doc, headings)
            
            # Combine scores
            return int((structure_score + format_score) / 2)
        except:
            return 50  # Default middle score if error
    
    def _check_heading_formatting(self, doc: Document, headings: List) -> int:
        """Check if headings have proper formatting."""
        try:
            if not headings:
                return 0
            
            # Get heading style requirements from style guide
            style_specs = self.style_guide.get("styles", {})
            
            # Count correctly formatted headings
            correct_format = 0
            
            for h in headings:
                heading_style = h.style.name
                
                # Check if style exists in style guide
                if heading_style in style_specs:
                    correct_format += 1
            
            return (correct_format / len(headings)) * 100
        except:
            return 50  # Default middle score if error
    
    def _score_fonts(self, doc: Document) -> int:
        """Score font compliance."""
        try:
            # Get required font from style guide
            required_font = self.style_guide.get("meta", {}).get("default_font", "Times New Roman")
            
            # Check paragraphs for proper font
            if not doc.paragraphs:
                return 50  # Default middle score if no paragraphs
            
            # Count paragraphs with correct font
            correct_font = 0
            total_checked = 0
            
            for p in doc.paragraphs:
                if p.text.strip():  # Only check non-empty paragraphs
                    total_checked += 1
                    
                    # Check if any run in paragraph has correct font
                    if any(run.font.name == required_font for run in p.runs if run.font.name):
                        correct_font += 1
            
            if total_checked == 0:
                return 50  # Default middle score if no valid paragraphs
            
            # Calculate score
            score = (correct_font / total_checked) * 100
            return int(score)
        except:
            return 50  # Default middle score if error
    
    def _score_paragraphs(self, doc: Document) -> int:
        """Score paragraph formatting compliance."""
        try:
            # Check paragraphs for proper first-line indentation and alignment
            if not doc.paragraphs:
                return 50  # Default middle score if no paragraphs
            
            # Get normal paragraph style from style guide
            normal_style = self.style_guide.get("styles", {}).get("Normal", {})
            required_indent = normal_style.get("paragraph", {}).get("first_line_indent", Inches(0.5))
            required_alignment = normal_style.get("paragraph", {}).get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
            
            # Count paragraphs with correct formatting
            correct_format = 0
            total_checked = 0
            
            for p in doc.paragraphs:
                if p.text.strip() and not p.style.name.startswith('Heading'):
                    total_checked += 1
                    
                    # Check indentation and alignment
                    try:
                        indent_correct = (abs(p.paragraph_format.first_line_indent.inches - required_indent.inches) <= 0.1)
                        alignment_correct = (p.paragraph_format.alignment == required_alignment)
                        
                        if indent_correct and alignment_correct:
                            correct_format += 1
                    except:
                        pass  # Skip if cannot determine formatting
            
            if total_checked == 0:
                return 50  # Default middle score if no valid paragraphs
            
            # Calculate score
            score = (correct_format / total_checked) * 100
            return int(score)
        except:
            return 50  # Default middle score if error
    
    def _score_references(self, doc: Document) -> int:
        """Score reference formatting compliance."""
        try:
            # Find reference section
            references_start = -1
            
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().lower() == "references":
                    references_start = i
                    break
            
            if references_start == -1 or references_start >= len(doc.paragraphs) - 1:
                return 40  # Penalize for missing or empty references
            
            # Check reference entries
            reference_entries = []
            for i in range(references_start + 1, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip()
                if text and not text.lower() in ["appendix", "appendices"]:
                    reference_entries.append(doc.paragraphs[i])
                elif text.lower() in ["appendix", "appendices"]:
                    break  # Stop at appendix section
            
            if not reference_entries:
                return 40  # Penalize for no reference entries
            
            # Score based on formatting consistency and hanging indent
            consistent_format = 0
            hanging_indent = 0
            
            for ref in reference_entries:
                try:
                    # Check for hanging indent
                    if ref.paragraph_format.first_line_indent and ref.paragraph_format.left_indent:
                        if ref.paragraph_format.first_line_indent.inches < 0:
                            hanging_indent += 1
                    
                    # Check for consistent formatting
                    if self._check_reference_format(ref.text):
                        consistent_format += 1
                except:
                    pass
            
            # Calculate scores
            indent_score = (hanging_indent / len(reference_entries)) * 100
            format_score = (consistent_format / len(reference_entries)) * 100
            
            # Combine scores (weighted)
            return int((indent_score * 0.4) + (format_score * 0.6))
        except:
            return 50  # Default middle score if error
    
    def _check_reference_format(self, reference_text: str) -> bool:
        """Check if reference entry follows expected format."""
        # Basic check for APA style: Author. (Year). Title. Source.
        if self.style_name == "apa":
            return bool(re.search(r'\([0-9]{4}\)', reference_text))
        
        # Basic check for MLA style: Author. "Title." Source, Year.
        elif self.style_name == "mla":
            return bool(re.search(r'["\'].*["\']', reference_text) and re.search(r'[12][0-9]{3}', reference_text))
        
        # Basic check for Chicago style: Author. Year. "Title." Source.
        elif self.style_name == "chicago":
            return bool(re.search(r'[12][0-9]{3}', reference_text) and re.search(r'["\'].*["\']', reference_text))
        
        return True  # Default to true for unknown styles
    
    def _score_citations(self, doc: Document) -> int:
        """Score in-text citation compliance."""
        try:
            # Get expected citation format from style guide
            citation_format = self.style_guide.get("meta", {}).get("citation_format", "({Author}, {Year})")
            
            # Extract all in-text citations
            citations = []
            for p in doc.paragraphs:
                citations.extend(self._extract_citations(p.text))
            
            if not citations:
                return 40  # Penalize for no citations
            
            # Check citation format compliance
            compliant_citations = 0
            
            for citation in citations:
                if self._check_citation_format(citation):
                    compliant_citations += 1
            
            # Calculate score
            score = (compliant_citations / len(citations)) * 100
            return int(score)
        except:
            return 50  # Default middle score if error
    
    def _extract_citations(self, text: str) -> List[str]:
        """Extract potential citations from text."""
        citations = []
        
        # APA style citations: (Author, Year)
        if self.style_name == "apa":
            matches = re.finditer(r'\([^)]+\d{4}[^)]*\)', text)
            citations.extend(match.group(0) for match in matches)
        
        # MLA style citations: (Author Page)
        elif self.style_name == "mla":
            matches = re.finditer(r'\([^)]+\d+[^)]*\)', text)
            citations.extend(match.group(0) for match in matches)
        
        # Chicago style citations: (Author Year, Page)
        elif self.style_name == "chicago":
            matches = re.finditer(r'\([^)]+\d{4}[^)]*\)', text)
            citations.extend(match.group(0) for match in matches)
        
        # General parenthetical citations
        else:
            matches = re.finditer(r'\([^)]+\)', text)
            citations.extend(match.group(0) for match in matches)
        
        return citations
    
    def _check_citation_format(self, citation: str) -> bool:
        """Check if citation matches expected format."""
        # APA style: (Author, Year) or (Author, Year, p. X)
        if self.style_name == "apa":
            return bool(re.match(r'\([A-Za-z]+,\s+\d{4}', citation))
        
        # MLA style: (Author Page) or (Author XX)
        elif self.style_name == "mla":
            return bool(re.match(r'\([A-Za-z]+\s+\d+', citation))
        
        # Chicago style: (Author Year, Page) or (Author Year)
        elif self.style_name == "chicago":
            return bool(re.match(r'\([A-Za-z]+\s+\d{4}', citation))
        
        return True  # Default to true for unknown styles
    
    def _score_title_page(self, doc: Document) -> int:
        """Score title page compliance."""
        try:
            # Check if style requires title page
            requires_title_page = self.style_guide.get("meta", {}).get("title_page", False)
            has_title_page = self._has_title_page(doc)
            
            # If title page is required but missing, penalize heavily
            if requires_title_page and not has_title_page:
                return 20
            
            # If title page is not required but present, minor penalty
            if not requires_title_page and has_title_page:
                return 80
            
            # If title page requirement matches actual document, check formatting
            if has_title_page:
                # Check title page formatting
                title_format_score = self._check_title_page_formatting(doc)
                return title_format_score
            else:
                return 100  # No title page required and none present
        except:
            return 50  # Default middle score if error
    
    def _check_title_page_formatting(self, doc: Document) -> int:
        """Check if title page has proper formatting."""
        try:
            # Check first few paragraphs for title page elements
            if len(doc.paragraphs) < 3:
                return 50  # Not enough paragraphs to assess
            
            # For APA style
            if self.style_name == "apa":
                # Check for centered title
                if doc.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return 90
                else:
                    return 60
            
            # For MLA style (typically no separate title page)
            elif self.style_name == "mla":
                return 70  # MLA typically doesn't use title page
            
            # For Chicago style
            elif self.style_name == "chicago":
                # Check for centered title
                if doc.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return 90
                else:
                    return 60
            
            return 70  # Default score for other styles
        except:
            return 50  # Default middle score if error
    
    def _find_formatting_issues(self, doc: Document) -> List[Dict]:
        """Find specific formatting issues in the document."""
        issues = []
        
        # Check margins
        margin_issues = self._check_margin_issues(doc)
        issues.extend(margin_issues)
        
        # Check line spacing
        spacing_issues = self._check_spacing_issues(doc)
        issues.extend(spacing_issues)
        
        # Check heading structure
        heading_issues = self._check_heading_issues(doc)
        issues.extend(heading_issues)
        
        # Check font usage
        font_issues = self._check_font_issues(doc)
        issues.extend(font_issues)
        
        # Check reference formatting
        reference_issues = self._check_reference_issues(doc)
        issues.extend(reference_issues)
        
        # Check citation formatting
        citation_issues = self._check_citation_issues(doc)
        issues.extend(citation_issues)
        
        # Sort issues by severity
        return sorted(issues, key=lambda x: x["severity"], reverse=True)
    
    def _check_margin_issues(self, doc: Document) -> List[Dict]:
        """Check for margin-related issues."""
        issues = []
        
        try:
            if not doc.sections:
                return []
            
            # Get required margins from style guide
            required_margins = self.style_guide.get("margins", {})
            
            # Get actual margins from first section
            section = doc.sections[0]
            
            # Check each margin type
            margin_types = {
                "left": ("left_margin", "Left"),
                "right": ("right_margin", "Right"),
                "top": ("top_margin", "Top"),
                "bottom": ("bottom_margin", "Bottom")
            }
            
            for margin_key, (attr_name, display_name) in margin_types.items():
                if margin_key in required_margins:
                    required = required_margins[margin_key]
                    actual = getattr(section, attr_name)
                    
                    # Check if margin is significantly different
                    if abs(actual.inches - required.inches) > 0.1:
                        issues.append({
                            "category": "margins",
                            "issue": f"{display_name} margin incorrect",
                            "details": f"Current: {actual.inches:.2f}\", Required: {required.inches:.2f}\"",
                            "severity": "high" if abs(actual.inches - required.inches) > 0.3 else "medium",
                            "location": "Document sections"
                        })
        except Exception as e:
            issues.append({
                "category": "margins",
                "issue": "Could not check margins",
                "details": str(e),
                "severity": "low",
                "location": "Document sections"
            })
        
        return issues
    
    def _check_spacing_issues(self, doc: Document) -> List[Dict]:
        """Check for line spacing issues."""
        issues = []
        
        try:
            # Most academic styles require double spacing
            required_spacing = WD_LINE_SPACING.DOUBLE
            
            # Check a sample of paragraphs
            spacing_issues_count = 0
            paragraphs_checked = 0
            
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip() and not p.style.name.startswith('Heading'):
                    paragraphs_checked += 1
                    
                    try:
                        if p.paragraph_format.line_spacing_rule != required_spacing:
                            spacing_issues_count += 1
                            
                            # Only report specific instances for first few issues
                            if spacing_issues_count <= 3:
                                issues.append({
                                    "category": "spacing",
                                    "issue": "Incorrect line spacing",
                                    "details": f"Paragraph should use double spacing",
                                    "severity": "medium",
                                    "location": f"Paragraph {i+1}: \"{p.text[:40]}...\""
                                })
                    except:
                        pass
            
            # Add summary issue if many spacing problems found
            if spacing_issues_count > 3:
                issues.append({
                    "category": "spacing",
                    "issue": "Widespread spacing issues",
                    "details": f"Found {spacing_issues_count} paragraphs with incorrect spacing",
                    "severity": "high" if spacing_issues_count > 10 else "medium",
                    "location": "Throughout document"
                })
        except Exception as e:
            issues.append({
                "category": "spacing",
                "issue": "Could not check line spacing",
                "details": str(e),
                "severity": "low",
                "location": "Document paragraphs"
            })
        
        return issues
    
    def _check_heading_issues(self, doc: Document) -> List[Dict]:
        """Check for heading structure issues."""
        issues = []
        
        try:
            # Find all headings
            headings = []
            for i, p in enumerate(doc.paragraphs):
                if p.style.name.startswith('Heading '):
                    try:
                        level = int(p.style.name.split(' ')[1])
                        headings.append((i, level, p.text))
                    except (ValueError, IndexError):
                        pass
            
            if not headings:
                issues.append({
                    "category": "headings",
                    "issue": "No headings found",
                    "details": "Document lacks properly formatted headings",
                    "severity": "medium",
                    "location": "Throughout document"
                })
                return issues
            
            # Check for skipped heading levels
            prev_level = headings[0][1]
            for i, (para_idx, level, text) in enumerate(headings[1:], 1):
                if level > prev_level + 1:
                    issues.append({
                        "category": "headings",
                        "issue": "Skipped heading level",
                        "details": f"Heading jumps from H{prev_level} to H{level}",
                        "severity": "high",
                        "location": f"Heading: \"{text}\""
                    })
                prev_level = level
            
            # Check heading formatting
            for para_idx, level, text in headings:
                # Check if heading is empty
                if not text.strip():
                    issues.append({
                        "category": "headings",
                        "issue": "Empty heading",
                        "details": f"Heading {level} has no content",
                        "severity": "high",
                        "location": f"Paragraph {para_idx+1}"
                    })
        except Exception as e:
            issues.append({
                "category": "headings",
                "issue": "Could not check heading structure",
                "details": str(e),
                "severity": "low",
                "location": "Document headings"
            })
        
        return issues
    
    def _check_font_issues(self, doc: Document) -> List[Dict]:
        """Check for font-related issues."""
        issues = []
        
        try:
            # Get required font from style guide
            required_font = self.style_guide.get("meta", {}).get("default_font", "Times New Roman")
            
            # Track font usage
            font_counts = {}
            inconsistent_paras = []
            
            for i, p in enumerate(doc.paragraphs):
                if not p.text.strip():
                    continue
                
                # Check fonts used in this paragraph
                para_fonts = set()
                for run in p.runs:
                    if run.font.name:
                        para_fonts.add(run.font.name)
                        if run.font.name not in font_counts:
                            font_counts[run.font.name] = 0
                        font_counts[run.font.name] += 1
                
                # Check if paragraph uses multiple fonts
                if len(para_fonts) > 1:
                    inconsistent_paras.append((i, p.text[:40], para_fonts))
            
            # Report use of non-required fonts
            non_standard_fonts = [font for font in font_counts.keys() if font != required_font]
            if non_standard_fonts:
                issues.append({
                    "category": "font",
                    "issue": "Non-standard fonts used",
                    "details": f"Found {', '.join(non_standard_fonts[:3])}{'...' if len(non_standard_fonts) > 3 else ''}",
                    "severity": "high" if font_counts.get(required_font, 0) < sum(font_counts.values()) / 2 else "medium",
                    "location": "Throughout document"
                })
            
            # Report inconsistent font usage within paragraphs
            if inconsistent_paras:
                for i, text, fonts in inconsistent_paras[:3]:  # Report first 3
                    issues.append({
                        "category": "font",
                        "issue": "Inconsistent fonts in paragraph",
                        "details": f"Paragraph uses multiple fonts: {', '.join(fonts)}",
                        "severity": "medium",
                        "location": f"Paragraph {i+1}: \"{text}...\""
                    })
        except Exception as e:
            issues.append({
                "category": "font",
                "issue": "Could not check fonts",
                "details": str(e),
                "severity": "low",
                "location": "Document text"
            })
        
        return issues
    
    def _check_reference_issues(self, doc: Document) -> List[Dict]:
        """Check for reference-related issues."""
        issues = []
        
        try:
            # Find reference section
            references_start = -1
            
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().lower() == "references":
                    references_start = i
                    break
            
            if references_start == -1:
                issues.append({
                    "category": "references",
                    "issue": "No references section found",
                    "details": "Document lacks a properly labeled References section",
                    "severity": "high",
                    "location": "End of document"
                })
                return issues
            
            # Check reference entries
            reference_entries = []
            for i in range(references_start + 1, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip()
                if text and not text.lower() in ["appendix", "appendices"]:
                    reference_entries.append((i, text))
                elif text.lower() in ["appendix", "appendices"]:
                    break
            
            if not reference_entries:
                issues.append({
                    "category": "references",
                    "issue": "Empty references section",
                    "details": "References section contains no entries",
                    "severity": "high",
                    "location": f"After paragraph {references_start+1}"
                })
                return issues
            
            # Check reference formatting
            hanging_indent_missing = []
            format_issues = []
            
            for i, text in reference_entries:
                # Check for hanging indent
                try:
                    p = doc.paragraphs[i]
                    if not (p.paragraph_format.first_line_indent and 
                           p.paragraph_format.first_line_indent.inches < 0 and
                           p.paragraph_format.left_indent and 
                           p.paragraph_format.left_indent.inches > 0):
                        hanging_indent_missing.append(i)
                except:
                    pass
                
                # Check reference format
                if not self._check_reference_format(text):
                    format_issues.append((i, text[:40]))
            
            # Report hanging indent issues
            if hanging_indent_missing and len(hanging_indent_missing) > len(reference_entries) / 2:
                issues.append({
                    "category": "references",
                    "issue": "Missing hanging indentation",
                    "details": "References should use hanging indentation",
                    "severity": "high",
                    "location": "References section"
                })
            
            # Report format issues
            if format_issues:
                issues.append({
                    "category": "references",
                    "issue": f"Incorrect reference format ({self.style_name.upper()})",
                    "details": f"Found {len(format_issues)} references with formatting issues",
                    "severity": "high",
                    "location": "References section"
                })
                
                # Report specific examples
                for i, text in format_issues[:3]:
                    issues.append({
                        "category": "references",
                        "issue": "Reference format issue",
                        "details": f"Entry doesn't match {self.style_name.upper()} format",
                        "severity": "medium",
                        "location": f"Reference: \"{text}...\""
                    })
        except Exception as e:
            issues.append({
                "category": "references",
                "issue": "Could not check references",
                "details": str(e),
                "severity": "low",
                "location": "References section"
            })
        
        return issues
    
    def _check_citation_issues(self, doc: Document) -> List[Dict]:
        """Check for citation-related issues."""
        issues = []
        
        try:
            # Extract all citations
            all_citations = []
            for i, p in enumerate(doc.paragraphs):
                text = p.text
                citations = self._extract_citations(text)
                for citation in citations:
                    all_citations.append((i, citation, text[:40]))
            
            if not all_citations:
                issues.append({
                    "category": "citations",
                    "issue": "No in-text citations found",
                    "details": f"Document lacks in-text citations in {self.style_name.upper()} format",
                    "severity": "high",
                    "location": "Throughout document"
                })
                return issues
            
            # Check citation format
            format_issues = []
            for i, citation, context in all_citations:
                if not self._check_citation_format(citation):
                    format_issues.append((i, citation, context))
            
            # Report format issues
            if format_issues and len(format_issues) > len(all_citations) / 3:
                issues.append({
                    "category": "citations",
                    "issue": f"Widespread citation format issues ({self.style_name.upper()})",
                    "details": f"Found {len(format_issues)} citations with incorrect format",
                    "severity": "high",
                    "location": "Throughout document"
                })
                
                # Report specific examples
                for i, citation, context in format_issues[:3]:
                    issues.append({
                        "category": "citations",
                        "issue": "Citation format issue",
                        "details": f"{citation} doesn't match {self.style_name.upper()} format",
                        "severity": "medium",
                        "location": f"Paragraph {i+1}: \"{context}...\""
                    })
        except Exception as e:
            issues.append({
                "category": "citations",
                "issue": "Could not check citations",
                "details": str(e),
                "severity": "low",
                "location": "Document text"
            })
        
        return issues
    
    def _analyze_citations(self, doc: Document) -> Dict:
        """Analyze in-text citations in the document."""
        try:
            # Extract all citations
            citations = []
            for i, p in enumerate(doc.paragraphs):
                citations.extend(self._extract_citations(p.text))
            
            # Count unique citation patterns
            citation_counts = {}
            for citation in citations:
                normalized = citation.lower()
                if normalized not in citation_counts:
                    citation_counts[normalized] = 0
                citation_counts[normalized] += 1
            
            # Calculate statistics
            return {
                "total_citations": len(citations),
                "unique_citations": len(citation_counts),
                "most_common": sorted(citation_counts.items(), key=lambda x: x[1], reverse=True)[:5],
                "citation_density": len(citations) / len(doc.paragraphs) if doc.paragraphs else 0,
                "format_compliance": sum(1 for c in citations if self._check_citation_format(c)) / len(citations) if citations else 0
            }
        except:
            return {
                "total_citations": 0,
                "unique_citations": 0,
                "most_common": [],
                "citation_density": 0,
                "format_compliance": 0
            }
    
    def _analyze_references(self, doc: Document) -> Dict:
        """Analyze reference list in the document."""
        try:
            # Find reference section
            references_start = -1
            
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().lower() == "references":
                    references_start = i
                    break
            
            if references_start == -1 or references_start >= len(doc.paragraphs) - 1:
                return {
                    "found": False,
                    "entries": 0,
                    "format_compliance": 0
                }
            
            # Extract reference entries
            reference_entries = []
            for i in range(references_start + 1, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip()
                if text and not text.lower() in ["appendix", "appendices"]:
                    reference_entries.append(text)
                elif text.lower() in ["appendix", "appendices"]:
                    break
            
            # Calculate statistics
            return {
                "found": True,
                "entries": len(reference_entries),
                "format_compliance": sum(1 for r in reference_entries if self._check_reference_format(r)) / len(reference_entries) if reference_entries else 0,
                "has_hanging_indent": self._check_references_have_hanging_indent(doc, references_start)
            }
        except:
            return {
                "found": False,
                "entries": 0,
                "format_compliance": 0
            }
    
    def _check_references_have_hanging_indent(self, doc: Document, references_start: int) -> bool:
        """Check if references use hanging indent."""
        try:
            hanging_indents = 0
            total_checked = 0
            
            for i in range(references_start + 1, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip()
                if not text:
                    continue
                if text.lower() in ["appendix", "appendices"]:
                    break
                
                # Check paragraph format
                total_checked += 1
                p = doc.paragraphs[i]
                
                try:
                    if (p.paragraph_format.first_line_indent and 
                        p.paragraph_format.first_line_indent.inches < 0 and
                        p.paragraph_format.left_indent and 
                        p.paragraph_format.left_indent.inches > 0):
                        hanging_indents += 1
                except:
                    pass
            
            return hanging_indents > total_checked / 2 if total_checked > 0 else False
        except:
            return False
    
    def _analyze_heading_structure(self, doc: Document) -> Dict:
        """Analyze document heading structure."""
        try:
            # Find all headings
            headings = []
            for i, p in enumerate(doc.paragraphs):
                if p.style.name.startswith('Heading '):
                    try:
                        level = int(p.style.name.split(' ')[1])
                        headings.append({
                            "level": level,
                            "text": p.text,
                            "paragraph": i
                        })
                    except (ValueError, IndexError):
                        pass
            
            # Analyze level distribution
            level_counts = {}
            for heading in headings:
                level = heading["level"]
                if level not in level_counts:
                    level_counts[level] = 0
                level_counts[level] += 1
            
            # Check for logical structure
            structure_issues = 0
            prev_level = headings[0]["level"] if headings else 0
            for heading in headings[1:]:
                if heading["level"] > prev_level + 1:
                    structure_issues += 1
                prev_level = heading["level"]
            
            return {
                "total_headings": len(headings),
                "by_level": level_counts,
                "structure_issues": structure_issues,
                "headings": headings[:10]  # Return first 10 headings
            }
        except:
            return {
                "total_headings": 0,
                "by_level": {},
                "structure_issues": 0,
                "headings": []
            }
    
    def _analyze_spacing(self, doc: Document) -> Dict:
        """Analyze document line spacing."""
        try:
            # Count paragraphs with different spacing
            spacing_counts = {
                "single": 0,
                "1.5": 0,
                "double": 0,
                "other": 0
            }
            
            total_checked = 0
            
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                
                try:
                    spacing = p.paragraph_format.line_spacing_rule
                    if spacing == WD_LINE_SPACING.SINGLE:
                        spacing_counts["single"] += 1
                    elif spacing == 1.5:
                        spacing_counts["1.5"] += 1
                    elif spacing == WD_LINE_SPACING.DOUBLE:
                        spacing_counts["double"] += 1
                    else:
                        spacing_counts["other"] += 1
                    
                    total_checked += 1
                except:
                    pass
            
            # Calculate predominant spacing
            predominant = "unknown"
            max_count = 0
            for spacing, count in spacing_counts.items():
                if count > max_count:
                    max_count = count
                    predominant = spacing
            
            return {
                "counts": spacing_counts,
                "predominant": predominant,
                "compliance": spacing_counts["double"] / total_checked if total_checked > 0 else 0
            }
        except:
            return {
                "counts": {"single": 0, "1.5": 0, "double": 0, "other": 0},
                "predominant": "unknown",
                "compliance": 0
            }
    
    def _analyze_margins(self, doc: Document) -> Dict:
        """Analyze document margins."""
        try:
            if not doc.sections:
                return {
                    "detected": False
                }
            
            # Get margins from first section
            section = doc.sections[0]
            
            # Get required margins from style guide
            required_margins = self.style_guide.get("margins", {})
            
            margins = {
                "left": section.left_margin.inches,
                "right": section.right_margin.inches,
                "top": section.top_margin.inches,
                "bottom": section.bottom_margin.inches
            }
            
            # Check compliance
            compliance = {}
            for margin_type, actual in margins.items():
                required = required_margins.get(margin_type, Inches(1)).inches
                compliance[margin_type] = abs(actual - required) <= 0.1
            
            return {
                "detected": True,
                "margins": margins,
                "required": {k: v.inches for k, v in required_margins.items() if hasattr(v, "inches")},
                "compliance": compliance
            }
        except:
            return {
                "detected": False
            }
    
    def generate_report(self, analysis_result: Dict) -> str:
        """
        Generate a human-readable formatting analysis report.
        
        Args:
            analysis_result: Result from analyze_document()
            
        Returns:
            Formatted string report
        """
        lines = []
        lines.append("=" * 60)
        lines.append(f"  FORMATLY V3 - {self.style_name.upper()} FORMATTING ANALYSIS")
        lines.append("=" * 60)
        lines.append("")
        
        # Document info
        doc_info = analysis_result.get("document_info", {})
        lines.append("📄 DOCUMENT INFORMATION")
        lines.append(f"  Style Guide: {doc_info.get('style', self.style_name.upper())}")
        lines.append(f"  Word Count: {doc_info.get('word_count', 0)}")
        lines.append(f"  Paragraphs: {doc_info.get('paragraphs', 0)}")
        lines.append("")
        
        # Compliance score
        score = analysis_result.get("compliance_score", {})
        lines.append("📊 FORMATTING COMPLIANCE")
        lines.append(f"  Overall Score: {score.get('overall', 0)}/100 ({score.get('grade', 'F')})")
        lines.append("  Category Scores:")
        
        for category, cat_score in score.get("details", {}).items():
            lines.append(f"    - {category.title()}: {cat_score}/100")
        lines.append("")
        
        # Critical issues
        issues = analysis_result.get("formatting_issues", [])
        high_issues = [i for i in issues if i.get("severity") == "high"]
        
        if high_issues:
            lines.append("❗ CRITICAL FORMATTING ISSUES")
            for issue in high_issues[:5]:  # Show top 5
                lines.append(f"  • {issue.get('issue')}")
                lines.append(f"    {issue.get('details')}")
                lines.append(f"    Location: {issue.get('location')}")
            
            if len(high_issues) > 5:
                lines.append(f"  ... and {len(high_issues) - 5} more critical issues")
            lines.append("")
        
        # Citations
        citation_analysis = analysis_result.get("citation_analysis", {})
        lines.append("🔍 CITATION ANALYSIS")
        lines.append(f"  Total In-text Citations: {citation_analysis.get('total_citations', 0)}")
        lines.append(f"  Format Compliance: {citation_analysis.get('format_compliance', 0) * 100:.1f}%")
        lines.append("")
        
        # References
        ref_analysis = analysis_result.get("reference_analysis", {})
        lines.append("📚 REFERENCES ANALYSIS")
        if ref_analysis.get("found", False):
            lines.append(f"  Reference Entries: {ref_analysis.get('entries', 0)}")
            lines.append(f"  Format Compliance: {ref_analysis.get('format_compliance', 0) * 100:.1f}%")
            lines.append(f"  Hanging Indent: {'Yes' if ref_analysis.get('has_hanging_indent', False) else 'No'}")
        else:
            lines.append("  No references section found")
        lines.append("")
        
        # Margins
        margin_analysis = analysis_result.get("margin_analysis", {})
        if margin_analysis.get("detected", False):
            lines.append("📏 MARGINS")
            margins = margin_analysis.get("margins", {})
            for margin, value in margins.items():
                compliance = "✓" if margin_analysis.get("compliance", {}).get(margin, False) else "✗"
                lines.append(f"  {margin.title()}: {value:.2f}\" {compliance}")
            lines.append("")
        
        # Summary of recommendations
        lines.append("🔧 RECOMMENDED ACTIONS")
        
        # Generate recommendations based on issues
        recommendations = self._generate_recommendations(analysis_result)
        for rec in recommendations:
            lines.append(f"  • {rec}")
        
        lines.append("")
        lines.append("=" * 60)
        
        return "\n".join(lines)
    
    def _generate_recommendations(self, analysis_result: Dict) -> List[str]:
        """Generate formatting recommendations based on analysis."""
        recommendations = []
        issues = analysis_result.get("formatting_issues", [])
        
        # Group issues by category
        issues_by_category = {}
        for issue in issues:
            category = issue.get("category", "other")
            if category not in issues_by_category:
                issues_by_category[category] = []
            issues_by_category[category].append(issue)
        
        # Generate category-specific recommendations
        if "margins" in issues_by_category:
            recommendations.append(f"Adjust document margins to match {self.style_name.upper()} requirements")
        
        if "spacing" in issues_by_category:
            recommendations.append("Set line spacing to double throughout document")
        
        if "headings" in issues_by_category:
            recommendations.append("Fix heading structure - use proper heading levels without skipping")
        
        if "font" in issues_by_category:
            font = self.style_guide.get("meta", {}).get("default_font", "Times New Roman")
            recommendations.append(f"Use {font} font consistently throughout document")
        
        if "references" in issues_by_category:
            recommendations.append(f"Format references according to {self.style_name.upper()} style with hanging indent")
        
        if "citations" in issues_by_category:
            recommendations.append(f"Ensure all in-text citations follow {self.style_name.upper()} format")
        
        # Add general recommendation if needed
        if not recommendations:
            recommendations.append(f"Document generally follows {self.style_name.upper()} guidelines")
        
        # Add recommendation to use Formatly
        recommendations.append("Use Formatly's formatting tool to automatically correct these issues")
        
        return recommendations
