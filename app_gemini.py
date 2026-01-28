from core.api_clients import GeminiClient
from core.formatter import AdvancedFormatter
from style_guides import STYLE_GUIDES
from spell_check import DocumentChecker, format_error_report
from utils.auto_corrector import AutoCorrector
from utils.formatting_analyzer import FormattingAnalyzer
import os
import argparse
import winsound
import traceback
from pathlib import Path
from docx import Document

def list_available_styles():
    """List all available citation styles."""
    print("Available citation styles:")
    for style_name, style_config in STYLE_GUIDES.items():
        meta = style_config.get("meta", {})
        print(f"  {style_name.upper():8} - {meta.get('default_font', 'Times New Roman')} font")
        print(f"          {'Title page: ' + ('Yes' if meta.get('title_page') else 'No'):20}")
        print(f"          {'Abstract: ' + ('Required' if meta.get('requires_abstract') else 'Optional'):20}")
        print()

def validate_style(style_name: str) -> str:
    """Validate and return the style name, defaulting to APA if invalid."""
    return style_name if style_name.lower() in STYLE_GUIDES else "apa"

def determine_output_path(input_path: Path, output_arg: str, style_name: str) -> Path:
    """Determine the output path based on input and arguments."""
    if output_arg:
        return Path(output_arg)
    
    formatting_dir = Path("formatted")
    formatting_dir.mkdir(exist_ok=True)

    output_path = formatting_dir / f"{input_path.stem}_formatted_{style_name}{input_path.suffix}"
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")

    return output_path

def track_changes(input_path: Path, output_path: Path):
    """Track changes in document."""
    from utils.track_changes import TrackChanges
    
    # Convert to absolute paths
    input_file = str(input_path.absolute())
    output_file = str(output_path.absolute())
    
    print(f"Original file: {input_file}")
    print(f"Formatted file: {output_file}")
    
    tracker = TrackChanges(input_file, output_file)
    tracker.compare_docs()

def get_input_file_interactive():
    print("\nNote: If you provide just a filename, it will be searched in the 'documents' folder.")
    while True:
        input_file = input("Enter the path to your Word document: ").strip()
        if not input_file:
            print("Please provide a file path.")
            continue
        
        input_path = Path(input_file)
        
        if input_path.exists():
            if input_path.suffix.lower() == '.docx':
                return input_path
            else:
                print("Error: File must be a Word document (.docx)")
        else:
            documents_path = Path("documents") / input_path.name
            if documents_path.exists():
                if documents_path.suffix.lower() == '.docx':
                    return documents_path
                else:
                    print("Error: File must be a Word document (.docx)")
            else:
                print(f"Error: File '{input_file}' not found (also checked in 'documents' folder).")

def validate_input_file(input_path: str) -> Path:
    if not input_path:
        raise ValueError("Input file path is required")
    
    path = Path(input_path)
    
    if not path.exists():
        documents_path = Path("documents") / path.name
        if documents_path.exists():
            path = documents_path
        else:
            raise FileNotFoundError(f"Input file not found: {input_path} (also checked in 'documents' folder)")
    
    if path.suffix.lower() != '.docx':
        raise ValueError(f"Input file must be a Word document (.docx), got: {path.suffix}")
    
    return path

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Format documents according to citation styles using Gemini backend.')
    
    parser.add_argument('input', nargs='?', help='Input Word document path')
    
    parser.add_argument('-s', '--style', default='apa',
                      help='Citation style (default: apa). Use --list-styles to see available options.')
    parser.add_argument('-o', '--output', help='Output file path (default: input_filename_formatted_style.docx)')
    parser.add_argument('--english', choices=['us', 'gb', 'au', 'ca'], default='us',
                      help='English language variant for spelling (default: us)')
    
    parser.add_argument('-i', '--interactive', action='store_true',
                      help='Interactive mode - prompts for input file if not provided')
    parser.add_argument('-l', '--list-styles', action='store_true',
                      help='List available citation styles and exit')
    parser.add_argument('-t', '--track-changes', action='store_true',
                        help='Enable track changes in document')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--report-only', action='store_true', 
                      help='Generate comprehensive reports (spelling, formatting) without modifying the document.')
    group.add_argument('--fix-errors', action='store_true', 
                      help='Automatically fix spelling, punctuation, and capitalization errors, then format the document.')
    group.add_argument('--format', action='store_true', 
                      help='Apply document formatting based on the chosen style, skipping error correction.')
    parser.add_argument('--overwrite', action='store_true', help='Overwrite output file if it exists')

    mla_group = parser.add_mutually_exclusive_group()
    mla_group.add_argument('--mla-title-page', action='store_true',
                          help='For MLA style, generate a title page (no page number on the title page).')
    mla_group.add_argument('--mla-heading', action='store_true',
                          help='For MLA style, use a first-page heading instead of a title page (page numbering starts from page 1).')

    return parser.parse_args()

def main():
    args = parse_arguments()
    
    if args.list_styles:
        list_available_styles()
        return 0
    
    input_path = None
    if args.input:
        try:
            input_path = validate_input_file(args.input)
        except (FileNotFoundError, ValueError) as e:
            print(f"Error: {e}")
            return 1
    elif args.interactive:
        try:
            input_path = get_input_file_interactive()
        except KeyboardInterrupt:
            print("\nOperation cancelled by user.")
            return 1
    else:
        print("Error: Input file is required. Use --interactive mode or provide a file path.")
        return 1
    
    style_name = validate_style(args.style)
    
    try:
        # Load Client
        client = GeminiClient()
        print(f"Using Backend: Google Gemini ({client.model_name})")

        doc = Document(str(input_path))
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
        
        if args.report_only:
            print(f"Generating comprehensive report for: {input_path}")
            checker = DocumentChecker(english_variation=args.english)
            formatter_analyzer = FormattingAnalyzer(style_name)
            try:
                report = checker.get_correction_report(paragraphs_text)
                print("\n=== SPELLING REPORT ===")
                print(format_error_report(report))
                
                print("\n=== FORMATTING ANALYSIS ===")
                analysis_result = formatter_analyzer.analyze_document(str(input_path))
                analysis_report = formatter_analyzer.generate_report(analysis_result)
                print(analysis_report)
                
                print("\nReport generated successfully. No changes made to document.")
                return 0
            except Exception as e:
                print(f"Error during report generation: {e}")
                return 1
            finally:
                checker.close()
        
        elif args.fix_errors:
            print(f"Auto-fixing spelling, punctuation, and capitalization errors: {input_path}")
            auto_corrector = AutoCorrector(english_variation=args.english)
            
            try:
                print("\nApplying basic automatic corrections...")
                corrected_paragraphs, correction_summary = auto_corrector.apply_all(paragraphs_text)
                
                idx = 0
                for p in doc.paragraphs:
                    if p.text.strip() and idx < len(corrected_paragraphs):
                        p.text = corrected_paragraphs[idx]
                        idx += 1
                
                if correction_summary['total_corrections'] > 0:
                    print(auto_corrector.generate_correction_explanation(correction_summary))
                    by_type = correction_summary.get("corrections_by_type", {})
                    print(f"✅ Successfully applied {correction_summary['total_corrections']} corrections:")
                    print(f"   🔤 Spelling: {by_type.get('spelling', 0)}")
                    print(f"   🔣 Punctuation: {by_type.get('punctuation', 0)}")
                    print(f"   🔠 Capitalization: {by_type.get('capitalization', 0)}")
                else:
                    print("✅ No issues found that needed correction.")
                    
                print("\n📝 Proceeding with document formatting...")
            except Exception as e:
                print(f"Error during auto-fix: {e}")
                return 1
            finally:
                if auto_corrector:
                    auto_corrector.close()
        
        elif args.format:
            print(f"🎨 Applying formatting only (skipping error correction): {input_path}")

        if not args.report_only:
            print(f"📝 Formatting document: {input_path}")
            print(f"🎨 Style: {style_name.upper()}, English: {args.english.capitalize()}")
            
            # Use Core Formatter with Gemini Client
            formatter = AdvancedFormatter(style_name, ai_client=client)
            
            if style_name == 'mla':
                setattr(formatter, 'mla_heading', bool(getattr(args, 'mla_heading', False)))
                setattr(formatter, 'mla_title_page', bool(getattr(args, 'mla_title_page', False)))
                
            output_path = determine_output_path(input_path, args.output, style_name)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if output_path.exists() and not args.overwrite:
                response = input(f"Output file '{output_path}' already exists. Overwrite? (y/N): ")
                if response.lower() not in ['y', 'yes']:
                    print("Operation cancelled.")
                    return 0
            
            formatter.format_document(str(input_path), str(output_path), doc=doc)


            if args.track_changes:
                print(f"✅ Document formatted with Tracked Changes successfully: {output_path}")
                winsound.PlaySound("Beep", winsound.SND_ALIAS)
                track_changes(input_path, output_path)
            else:
                print(f"✅ Document formatted successfully: {output_path}")
                winsound.PlaySound("Beep", winsound.SND_ALIAS)
                if os.name == 'nt':
                     os.startfile(output_path)
                
    except Exception as e:
        print(f"Error: Failed to format document: {e}")
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    raise SystemExit(main())
