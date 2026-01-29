"""
Formatly CLI Application
------------------------
This module serves as the Command Line Interface (CLI) entry point for Formatly.
It utilizes the core logic located in the root `core/` directory to perform document
formatting, spell checking, and structure analysis.

Usage:
    python app.py [input_file] [options]

Key Dependencies:
    - core.formatter: Shared formatting logic
    - style_guides: Citation style definitions
    - utils: Helper utilities for rate limiting and analysis
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from style_guides import STYLE_GUIDES
from spell_check import DocumentChecker, format_error_report
from utils.auto_corrector import AutoCorrector
from utils.rate_limit_manager import RateLimitManager
from utils.formatting_analyzer import FormattingAnalyzer
import google.generativeai as genai
import json
import os
import re
import traceback
from dotenv import load_dotenv
from pathlib import Path
import argparse
import json5
import winsound
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)

class APIManager:
    """Manages API configuration and model setup."""
    
    def __init__(self):
        load_dotenv()
        self.api_key = os.getenv("GEMINI_API_KEY")
        self.model_name = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")
        self.rate_limit_manager = RateLimitManager(self.model_name)
        self._model = None
        
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY environment variable is not set")
            
    def get_model(self, system_instruction=None):
        """Get or initialize Gemini model with optional system instruction."""

        # Always create a new instance when system instruction is provided
        genai.configure(api_key=self.api_key)
        return genai.GenerativeModel(
            model_name=self.model_name,
            system_instruction=system_instruction
        )

        
    def execute_with_rate_limit(self, func, *args, **kwargs):
        """Execute function with rate limiting."""
        return self.rate_limit_manager.execute_with_rate_limit(func, *args, **kwargs)

# Global API manager instance
api_manager = APIManager()

# Expose API key and model name for other modules
API_KEY = api_manager.api_key
MODEL_NAME = api_manager.model_name

# Style mappings constant
STYLE_MAPPINGS = {
    # Front Matter / Preliminary Pages
    "title": "Title",
    "author": "Title Author",
    "institution": "Title Affiliation",
    "course": "Title Course",
    "instructor": "Title Instructor",
    "due_date": "Title Date",
    "title_department": "Title Department",
    "title_byline": "Title Byline",
    "dedication": "Heading 1",
    "acknowledgments": "Heading 1", # Often Heading 1 in structure
    "preface": "Heading 1",
    "epigraph": "Epigraph",
    "keywords": "Keywords",

    # Key Section Headings (Centralized Mapping)
    "abstract_heading": "Heading 1", 
    "references_heading": "Heading 1",
    "appendix_heading": "Appendix Title",
    "appendices_heading": "Appendices Title",
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "heading_4": "Heading 4",
    "heading_5": "Heading 5",

    # Content Types
    "abstract_text": "Abstract",
    "body": "Normal",
    "reference_list_item": "References",
    "appendix_text": "Normal",
    "list_item_number": "List Number",
    "list_item_bullet": "List Bullet",
    "list_item_alphabet": "List Alphabet",
    "list_item": "List Item",
    "figure_caption": "Figure Caption",
    "table_caption": "Table Title",
    "table_note": "Table Note",
    "table_data": "Table Body",
    "footnote": "Footnote Text",
    "block_quote": "Block Quote",
    "figure_text": "Normal",
    "code_block": "Code Block",
}

def configure_gemini(api_key: str) -> genai.GenerativeModel:
    """Configure and return the Gemini model."""
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(MODEL_NAME)

def load_style_guide(style_name: str) -> dict:
    """Load the appropriate style guide based on the requested style name."""
    style_name = style_name.lower()
    return STYLE_GUIDES.get(style_name, STYLE_GUIDES["apa"])

class DocumentStructureManager:
    """Manages document structure initialization, validation, and prompt creation."""
    
    @staticmethod
    def get_default_structure() -> dict:
        """Returns the default document structure template."""
        return {
            "blocks": [
                {
                    "type": "Normal", 
                    "content": "",
                    "attributes": {}
                }
            ]
        }
    
    @staticmethod
    def create_detection_prompt(paragraphs: list, english_variant: str = "us") -> tuple[str, str]:
        """Creates system and user prompts for document structure detection."""
        # Use a hardcoded, populated example to guide the AI better
        example_structure = {
            "blocks": [
                {
                    "type": "heading_1",
                    "content": "Introduction to Formatting",
                    "attributes": {}
                },
                {
                    "type": "body",
                    "content": "Artificial Intelligence is a rapidly evolving field...",
                    "attributes": {}
                }, 
                {
                    "type": "list_item_bullet", 
                    "content": "Machine Learning", 
                    "attributes": {"list_id": 1}
                }
            ]
        }
        json_structure = json.dumps(example_structure, indent=4)
        print(f'json_structure: {json_structure}')
        
        system_prompt = f"""You are an expert academic document parser and formatting assistant. Your task is to analyze the provided document content and extract it into a structured JSON format using a Linear Block Schema.
        
        TARGET JSON STRUCTURE:
        {json_structure}
        
        ⛔ CRITICAL EXTRACTION RULES:
        1. **Verbatim Extraction**: You must extract the text EXACTLY as it appears in the input. Do NOT summarize, Do NOT paraphrase, Do NOT remove words, Do NOT remove spaces, or "correct" typos/headings. Even if a heading appears redundant, extract it fully.
        2. **Content integrity**: Do not omit any paragraphs, headings, or sections. Every piece of text from the input must find a home in the JSON structure.
        3. **Valid JSON**: Your output must be strictly valid JSON.
        4. **CHAPTER TITLES**: Strings like "CHAPTER ONE", "CHAPTER 1", "CHAPTER I" are ALWAYS "heading_1".
        5. **CHAPTER NAMES & SECTIONS**: Top-level numbered sections (e.g., "1. Introduction", "2. Literature Review") MUST be "heading_1". HEADINGS following "CHAPTER X" must be "heading_1". Only demote to "heading_2" if the number indicates a subsection (e.g., "1.1", "2.3", "A.1", etc.).
        6. **HEADING LENGTH**: Headings are titles, NOT paragraphs. If a line of text is longer than ~12 words, it is likely "body" text, NOT a heading.
        7. **LISTS**: Assign the correct type and add a "list_id" (integer) field INSIDE "attributes".
           - **Type Selection**:
             - Use `list_item_bullet` for **unordered** lists.
             - Use `list_item_number` for **ordered**.
             - Use `list_item_alphabet` for **ordered**.
           - **Grouping**: Items in the same list MUST share the same "list_id".
           - **Restarts**: When a new list starts (e.g., after some body text), use a NEW "list_id".

        EXAMPLE BLOCK FOR LIST ITEM:
        {{
            "type": "list_item_number",
            "content": "First item",
            "attributes": {{"list_id": 1}}
        }}

        8. **No Empty Blocks**: Do NOT create blocks with empty content (content: ""). If a line is empty, ignore it.
        
        9. **APPENDICES**: 
           - Use `appendices_heading` for the master title "APPENDICES" or "LIST OF APPENDICES" (e.g., when it stands alone with sub appendices headings, eg. "Appendix A", "Appendix 1: Data Tables", "Conclusion of Appendix, etc.").
           - Use `appendix_heading` for individual appendix titles (e.g., "Appendix A", "Appendix 1: Data Tables", "Conclusion of Appendix, etc.").

        10. **Block Types**: Assign one of the following types to each block of content:
           - "title", "title_byline", "author", "institution", "title_department", "course", "instructor", "due_date" (Title Page Elements)
           - "keywords", "dedication", "epigraph", "acknowledgments", "preface" (Front Matter)
           - "heading_1", "heading_2", "heading_3", "heading_4", "heading_5". 
           - "abstract_heading", "abstract_text"
           - "body" (Standard Paragraphs)
           - "list_item_number", "list_item_bullet", "list_item_alphabet"
           - "block_quote"
           - "figure_caption", "table_caption", "table_note", "table_data"
           - "references_heading", "reference_list_item"
           - "appendices_heading", "appendix_heading"
           - "code_block", "footnote"

        11. **Headings for prelimnary pages (title page, abstract page, etc.) should be "heading_1"**

        Analyze the input document and populate the "blocks" list accordingly, intelligently."""
        
        user_prompt = f"Input content:\n\n{'\n'.join(paragraphs)}\n\nReturn JSON only."
        
        return system_prompt, user_prompt
    
    @staticmethod
    def parse_and_validate_response(response_text: str) -> dict:
        """Parses and validates AI response into document structure."""
        text = response_text.strip()
        
        # Clean markdown
        for marker in ["```json", "```"]:
            if text.startswith(marker):
                text = text[len(marker):]
            if text.endswith("```"):
                text = text[:-3]
        
        text = text.strip()
        
        # Validate basic JSON structure
        if not text or text[0] not in ['{', '[']:
            raise json.JSONDecodeError(f"Invalid JSON: {text[:100]}...", text, 0)
        
        try:
            data = json5.loads(text)
            return DocumentStructureManager.validate_structure(data)
        except (json.JSONDecodeError, ValueError) as e:
            print(f"JSON parse error: {e}")
            print(f"Raw response snippet: {response_text[:500]}...")
            raise
    
    @staticmethod
    def validate_structure(data: dict) -> dict:
        """Validates and sanitizes document structure."""
        if not isinstance(data, dict):
            raise ValueError("Document structure must be a dictionary")
        
        template = DocumentStructureManager.get_default_structure()
        
        # Deep merge with template
        def deep_merge(template_dict, data_dict):
            result = template_dict.copy()
            for key, value in data_dict.items():
                if key in result and isinstance(result[key], dict) and isinstance(value, dict):
                    result[key] = deep_merge(result[key], value)
                else:
                    result[key] = value
            return result
        
        return deep_merge(template, data)

def create_default_structure() -> dict:
    """Return a default document structure when parsing fails."""
    return {
        "blocks": []
    }

class AdvancedFormatter:
    def __init__(self, style_name):
        self.selected_style_name = style_name.lower()
        self.selected_style_guide = load_style_guide(self.selected_style_name)
        print(f"Loaded Style: {self.selected_style_name}")
        # Pre-compile reference formatting patterns
        self._compile_reference_patterns()
        # Compile inline formatting patterns
        self._compile_inline_patterns()
        # Cache for document styles
        self._doc_styles = None
        # Use the global API manager instead of having its own API key
        self.api_manager = api_manager

    def _compile_reference_patterns(self):
        """Pre-compile reference formatting patterns for efficiency."""
        ref_config = self.selected_style_guide.get('reference_formatting', {})
        self.compiled_ref_patterns = []
        for pattern in ref_config.get('patterns', []):
            if 'regex' in pattern and 'formatting' in pattern:
                self.compiled_ref_patterns.append({
                    'pattern': re.compile(pattern['regex'], re.DOTALL),
                    'formatting': pattern['formatting']
                })

    def _compile_inline_patterns(self):
        """Pre-compile inline formatting patterns."""
        inline_config = self.selected_style_guide.get('inline_formatting', [])
        self.compiled_inline_patterns = []
        for rule in inline_config:
            if 'regex' in rule and 'formatting' in rule:
                self.compiled_inline_patterns.append({
                    'description': rule.get('description', ''),
                    'target_styles': rule.get('target_styles', []),
                    'pattern': re.compile(rule['regex']), 
                    'formatting': rule['formatting']
                })
                
    def _apply_inline_formatting(self, doc, doc_structure):
        """
        Apply inline formatting based on regex patterns defined in style guide.
        """
        if not hasattr(self, 'compiled_inline_patterns') or not self.compiled_inline_patterns:
            return
        
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = p.style.name if hasattr(p.style, 'name') else str(p.style)
            
            for rule in self.compiled_inline_patterns:
                if rule['target_styles'] and style_name not in rule['target_styles']:
                    continue
                for match in rule['pattern'].finditer(p.text):
                    start_idx, end_idx = match.span()
                    self._apply_formatting_to_range(p, start_idx, end_idx, rule['formatting'])

    def _apply_formatting_to_range(self, paragraph, start_char, end_char, formatting):
        self._ensure_run_boundary(paragraph, start_char)
        self._ensure_run_boundary(paragraph, end_char)
        current_pos = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_end = current_pos + run_len
            if current_pos >= start_char and run_end <= end_char:
                if not run.font: run.font = paragraph.style.font 
                if 'italic' in formatting: run.font.italic = formatting['italic']
                if 'bold' in formatting: run.font.bold = formatting['bold']
            current_pos = run_end

    def _ensure_run_boundary(self, paragraph, char_index):
        current_pos = 0
        for i, run in enumerate(paragraph.runs):
            run_text = run.text
            run_len = len(run_text)
            if current_pos < char_index < current_pos + run_len:
                split_point = char_index - current_pos
                part1 = run_text[:split_point]
                part2 = run_text[split_point:]
                run.text = part1
                new_run = paragraph.add_run(part2)
                self._copy_run_formatting(run, new_run)
                p_element = paragraph._p
                run_element = run._element
                new_run_element = new_run._element
                p_element.remove(new_run_element)
                run_element.addnext(new_run_element)
                return
            current_pos += run_len

    def _get_style(self, doc, style_name):
        """Get an existing style from the document."""
        return doc.styles[style_name]

    def format_document(self, input_path, output_path):
        doc = Document(input_path)
        
        # Cache document styles
        self._doc_styles = doc.styles

        # Step 0: Remove leading whitespace
        self._remove_leading_whitespace(doc)

        # Step 0b: Clear old run formatting
        self._clear_run_formatting(doc)

        # Step 1: Apply base styles and margins
        self._customize_builtin_styles(doc)
        self._apply_margins(doc)

        # Step 2: Detect document structure
        paragraphs_text = [p.text for p in doc.paragraphs]
        doc_structure = self._detect_paragraph_types(paragraphs_text)

        # Step 3: Detect if a title page exists and find its boundary
        has_title_page, title_page_boundary = self._detect_and_manage_title_page(doc, doc_structure)

        # Step 4: Add page numbers based on title page presence
        # self._add_page_numbers(doc, has_title_page, title_page_boundary)

        # Step 5: Join split headings (e.g., CHAPTER X + Title)
        self._join_headings(doc, doc_structure)

        # Step 6: Format document content
        self._format_content_in_place(doc, doc_structure, has_title_page, title_page_boundary)

        # Step 7: Format references section
        self._format_references(doc, doc_structure)
        
        # Step 7b: Apply inline formatting (italics for Keywords, Note, etc.)
        self._apply_inline_formatting(doc, doc_structure)

        # Step 8: Apply consistent font properties
        self._apply_font_properties(doc)
        # Step 8b: Format tables
        self._format_tables(doc)

        self._apply_explicit_paragraph_properties(doc)
        # Step 7b: Apply explicit list/numbering properties (bullets / numbers)
        self._apply_list_properties(doc)
        self._apply_heading_styles(doc)

        # Step 8: Remove extra blank lines (but keep before title page boundary)
        self._remove_blank_lines(doc, title_page_boundary)


        # Final: Save document
        doc.save(output_path)

    def _customize_builtin_styles(self, doc):
        """
        Customize built-in styles and add custom styles based on the style guide.
        This method applies styles from style_guides.py to the document.
        """
        # Cache style configurations and document styles
        styles_config = self.selected_style_guide["styles"]
        doc_styles = doc.styles
        
        # Process each style in the configuration
        for style_name, style_config in styles_config.items():
            try:
                # Get the style. If it doesn't exist, create it.
                if style_name not in doc_styles:
                    style_type = style_config.get("type", WD_STYLE_TYPE.PARAGRAPH)
                    doc.styles.add_style(style_name, style_type)

                style = self._get_style(doc, style_name)
                
                # Set base style if specified
                if "based_on" in style_config and style_config["based_on"] in doc_styles:
                    style.base_style = doc_styles[style_config["based_on"]]
                
                # Apply font styles directly
                if "font" in style_config:
                    font_config = style_config["font"]
                    font = style.font

                    # Apply all font properties
                    if "name" in font_config:
                        font.name = font_config["name"]
                    if "size" in font_config:
                        font.size = font_config["size"]
                    if "bold" in font_config:
                        font.bold = font_config["bold"]
                    if "italic" in font_config:
                        font.italic = font_config["italic"]
                    if "underline" in font_config:
                        font.underline = font_config["underline"]
                    if "color" in font_config and font_config["color"]:
                        font.color.rgb = font_config["color"]
                    if "all_caps" in font_config:
                        font.all_caps = font_config["all_caps"]

                
                # Apply paragraph settings
                if "paragraph" in style_config:
                    para_format = style.paragraph_format
                    para_config = style_config["paragraph"]
                    
                    # Apply paragraph properties if they exist in the config
                    if "name" in para_config:
                        para_format.name = para_config["name"]
                    if "alignment" in para_config:
                        para_format.alignment = para_config["alignment"]
                    if "left_indent" in para_config:
                        para_format.left_indent = para_config["left_indent"]
                    if "right_indent" in para_config:
                        para_format.right_indent = para_config["right_indent"]
                    if "first_line_indent" in para_config:
                        para_format.first_line_indent = para_config["first_line_indent"]
                    if "space_before" in para_config:
                        para_format.space_before = para_config["space_before"]
                    if "space_after" in para_config:
                        para_format.space_after = para_config["space_after"]
                    if "line_spacing" in para_config:
                        para_format.line_spacing_rule = para_config["line_spacing"]
                    if "keep_together" in para_config:
                        para_format.keep_together = para_config["keep_together"]
                    if "keep_with_next" in para_config:
                        para_format.keep_with_next = para_config["keep_with_next"]
                    if "page_break_before" in para_config:
                        para_format.page_break_before = para_config["page_break_before"]
                    if "widow_control" in para_config:
                        para_format.widow_control = para_config["widow_control"]
                    if "orphan_control" in para_config:
                        try:
                            para_format.orphan_control = para_config["orphan_control"]
                        except AttributeError:
                            # orphan_control may not be available in all python-docx versions
                            pass
                    if "outline_level" in para_config:
                        try:
                            para_format.outline_level = para_config["outline_level"]
                        except AttributeError:
                            print("")
                            # outline_level may not be available in all python-docx versions
                            pass
                
                # Set common style properties
                if hasattr(style, 'hidden'):
                    style.hidden = style_config.get("hidden", False)
                if hasattr(style, 'unhide_when_used'):
                    style.unhide_when_used = style_config.get("unhide_when_used", True)
                if hasattr(style, 'quick_style'):
                    style.quick_style = style_config.get("quick_style", True)
                
                # Special handling for heading levels: ALWAYS enforce outline level
                if (style_name.startswith("Heading") and " " in style_name) or style_name == "Appendices Title":
                    try:
                        level = int(style_name.split(' ')[1]) if "Heading" in style_name else 1
                        # Force creation of pPr if needed and set outline level
                        if hasattr(style.paragraph_format, 'outline_level'):
                            style.paragraph_format.outline_level = level - 1
                        
                        # Fallback/Redundancy: Set it directly on new styles or if above failed
                        # This catches cases where creation didn't fully set properties
                        pPr = style._element.get_or_add_pPr()
                        outlineLvl = pPr.find(qn('w:outlineLvl'))
                        if outlineLvl is None:
                            outlineLvl = OxmlElement('w:outlineLvl')
                            pPr.append(outlineLvl)
                        outlineLvl.set(qn('w:val'), str(level - 1))
                        
                    except (ValueError, IndexError, AttributeError) as e:
                        print(f"Warning: Failed to set outline level for {style_name}: {e}")
                        pass
                
                # Set next style if specified
                if "next_style" in style_config and style_config["next_style"] in doc_styles:
                    style.next_paragraph_style = doc_styles[style_config["next_style"]]
                
                # Enforce no borders for academic styles (removes theme artifacts)
                self._remove_borders(style)
                
            except Exception as e:
                print(f"Warning: Could not apply style '{style_name}': {str(e)}")
                continue

    def _apply_margins(self, doc):
        section = doc.sections[0]
        margin_map = {
            "left": "left_margin", "right": "right_margin",
            "top": "top_margin", "bottom": "bottom_margin",
            "header": "header_distance", "footer": "footer_distance",
            "gutter": "gutter"
        }
        for key, value in self.selected_style_guide["margins"].items():
            if key in margin_map:
                setattr(section, margin_map[key], value)

    def _format_tables(self, doc):
        """
        Format all tables in the document.
        Ensures cells use the 'Table Content' style (or a 0-indent equivalent)
        to prevent text cutoff caused by the 'Normal' style's indent.
        """
        if "Table Content" not in doc.styles:
            return

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Avoid overwriting Heading styles if they exist in tables (rare but possible)
                        if paragraph.style.name.startswith("Heading"):
                             continue
                        
                        # Apply Table Content style
                        paragraph.style = doc.styles["Table Content"]

    def _remove_leading_whitespace(self, doc):
        """Remove leading whitespace (spaces, tabs) from the beginning of paragraphs."""
        for paragraph in doc.paragraphs:
            found_content = False
            for run in paragraph.runs:
                if found_content:
                    break
                
                text = run.text
                if not text:
                    continue
                
                if text.strip() == "":
                    # It's all whitespace (and we haven't found content yet), so clear it
                    run.text = ""
                else:
                    # It has some content. Strip the leading whitespace.
                    # This marks the start of content.
                    run.text = text.lstrip()
                    found_content = True

    def _clear_run_formatting(self, doc):
        """Clear old run formatting (character styles, fonts, colors, effects) from all paragraphs (body and tables)."""
        
        def clear_para_runs(paragraph):
            """Helper to clear runs in a single paragraph."""
            for run in paragraph.runs:
                # 1. Clear direct properties via python-docx
                run.bold = None
                run.italic = None
                run.underline = None
                
                if hasattr(run, 'font'):
                    run.font.name = None
                    run.font.size = None
                    try:
                        if run.font.color:
                            run.font.color.rgb = None
                    except (ValueError, AttributeError):
                        pass

                # 2. Robust XML Clearing: Remove persistent styling (rStyle, Color, Caps, etc.)
                rPr = run._element.rPr
                if rPr is not None:
                    # Tags to remove to ensure a completely clean slate
                    tags_to_remove = [
                        'w:rStyle',      # Character Style
                    ]
                    
                    for tag in tags_to_remove:
                        element = rPr.find(qn(tag))
                        if element is not None:
                            rPr.remove(element)

        # Clear Body Paragraphs
        for paragraph in doc.paragraphs:
            clear_para_runs(paragraph)
            
        # Clear Table Paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clear_para_runs(paragraph)

    def _remove_borders(self, style):
        """Removes all borders (bottom, top, left, right) from a paragraph style."""
        # Access the pPr (paragraph properties) element
        pPr = style._element.get_or_add_pPr()
        
        # Check for existing pBdr (paragraph borders) element
        pBdr = pPr.find(qn('w:pBdr'))
        
        # If it exists, remove it to clear all borders
        if pBdr is not None:
            pPr.remove(pBdr)

    def _detect_paragraph_types(self, paragraphs):

        system_prompt, user_prompt = DocumentStructureManager.create_detection_prompt(paragraphs)

        def _make_request(**kwargs):
            """Internal function to make the actual API request."""
            try:
                # Get model WITH system instruction
                model_with_system = self.api_manager.get_model(system_instruction=system_prompt)
                
                print("\n🤖 AI Response...\n")
                
                response = model_with_system.generate_content(
                    user_prompt,
                    request_options={'timeout': 600},  # 10 minutes timeout
                    stream=True
                    )
                
                # Stream and collect the response
                full_text = ""
                for chunk in response:
                    if chunk.text:
                        print(chunk.text, end='', flush=True)
                        full_text += chunk.text
                
                print("\n")  # New line after streaming
                
                # Validate response before parsing
                if not full_text:
                    raise ValueError("API returned an empty response. Please try again.")
                
                return DocumentStructureManager.parse_and_validate_response(full_text)
            except (json.JSONDecodeError, ValueError) as e:
                error_msg = f"JSON Parsing Error: {str(e)}"
                print(f"ERROR: {error_msg}")
                print(f"Response parsing failed. The API response may be incomplete or malformed.")
                raise ValueError(error_msg) from e
            except Exception as e:
                error_msg = f"API Request Error: {str(e)}"
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise ValueError(error_msg) from e

        return self.api_manager.execute_with_rate_limit(_make_request)

    def _detect_and_manage_title_page(self, doc, doc_structure):
        """
        No-op for Linear Block Schema.
        Returns False, 0 to treat everything as standard blocks.
        """
        return False, 0

    def _join_headings(self, doc, doc_structure):
        """
        Joins "CHAPTER X" and following title into a single paragraph.
        Modifies both doc and doc_structure.
        """
        blocks = doc_structure.get("blocks", [])
        if not blocks:
            return

        i = 0
        while i < len(blocks):
            block1 = blocks[i]
            content1 = block1.get("content", "").strip()
            type1 = block1.get("type", "")
            
            # Condition: First block is Heading 1 and starts with "CHAPTER"
            if type1 == "heading_1" and content1.upper().startswith("CHAPTER") and len(content1) < 20: 
                
                # Look ahead for the next content block
                next_content_idx = -1
                for j in range(i + 1, len(blocks)):
                    if blocks[j].get("content", "").strip():
                        next_content_idx = j
                        break
                
                if next_content_idx != -1:
                    block2 = blocks[next_content_idx]
                    content2 = block2.get("content", "").strip()
                    type2 = block2.get("type", "")
                    
                    # Merge condition: Next real content is Heading 1 or Heading 2
                    if type2 in ["heading_1"]:
                        # Attempt to find these in the document
                        p1 = None
                        p2 = None
                        
                        # We need to find them and ensure they are adjacent or close
                        # Scanning doc.paragraphs
                        for idx, p in enumerate(doc.paragraphs):
                            if p.text.strip() == content1:
                                p1 = p
                                # Check the next few paragraphs for content2 (ignore empty)
                                for offset in range(1, 10): # Widened search range
                                    if idx + offset < len(doc.paragraphs):
                                        next_p = doc.paragraphs[idx + offset]
                                        if next_p.text.strip() == content2:
                                            p2 = next_p
                                            break
                                        # Allow ignoring empty paragraphs between them
                                        if next_p.text.strip(): 
                                            break 
                                if p2:
                                    break
                        
                        if p1 and p2:
                            # Merge
                            separator = ": "
                            if content1[-1] in ":.-":
                                 separator = " "
                            
                            merged_text = content1 + separator + content2
                            p1.text = merged_text
                            
                            # Remove p2
                            p_element = p2._element
                            if p_element.getparent() is not None:
                                p_element.getparent().remove(p_element)
                            
                            # Update doc_structure
                            block1['content'] = merged_text
                            
                            # Remove block2 logic from structure
                            blocks.pop(next_content_idx)
                            
                            print(f"Merged headings: '{merged_text}'")
                            continue
            
            i += 1

    def _format_content_in_place(self, doc, doc_structure, has_title_page, title_page_boundary):
        """Format document content based on AI-detected structure."""
        style_map = {}

        # --- Build style_map from Linear Block Schema ---
        for block in doc_structure.get("blocks", []):
            content = block.get("content", "").strip()
            block_type = block.get("type", "Normal")
            # Handle list_id
            list_id = block.get("list_id")
            if list_id is None:
                list_id = block.get("attributes", {}).get("list_id")

            if content:
                style_map[content] = (STYLE_MAPPINGS.get(block_type, "Normal"), list_id)


        # --- Step 2: Iterate through existing paragraphs and apply styles ---
        # Create a list of (paragraph_object, original_text_stripped) tuples
        paragraphs_and_text = [(p, p.text.strip()) for p in doc.paragraphs]
        
        i = 0
        while i < len(paragraphs_and_text):
            p, current_text_stripped = paragraphs_and_text[i]



            
            if not current_text_stripped:
                i += 1
                continue # Skip empty paragraphs

            applied_style = False
            
            # 1. Check for exact matches
            if current_text_stripped in style_map:
                target_style_name, target_list_id = style_map[current_text_stripped]
                
                if target_style_name in doc.styles:
                    p.style = doc.styles[target_style_name]
                    applied_style = True
                else:
                    # Fallback for styles like "List Bullet" -> "List Bullet 1"
                    fallback_found = False
                    for suffix in [" 1", "1", " 2", "2"]:
                        variant_name = f"{target_style_name}{suffix}"
                        if variant_name in doc.styles:
                            p.style = doc.styles[variant_name]
                            target_style_name = variant_name
                            applied_style = True
                            fallback_found = True
                            print(f"Style '{target_style_name}' not found. Using variant '{variant_name}'.")
                            break
                    
                    if not fallback_found:
                        print(f"Warning: Required style '{target_style_name}' not found. Falling back to Normal.")
                        p.style = doc.styles["Normal"]
                        applied_style = True

                if applied_style and target_style_name != "Normal":
                    if target_list_id is not None:
                        try:
                            p.list_id = target_list_id
                        except:
                            pass
                    else:
                        # FIX: Explicitly remove list formatting (numPr) if we are switching to a non-list style
                        # This prevents bullet points from persisting when a user converts a list item to Reference or Heading
                        if p._element.pPr is not None and not (target_style_name.startswith("Heading") or "Appendix" in target_style_name): 
                            numPr = p._element.pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                p._element.pPr.remove(numPr)
                    

            # 2. Check for common section titles
            # Removed redundant hardcoded checks. We rely on the AI structure detection.

            # 3. Handle potential splits
            if not applied_style:
                potential_split_info = None
                for mapped_text, (mapped_style, mapped_list_id) in style_map.items():
                    if mapped_text and current_text_stripped.startswith(mapped_text) and len(current_text_stripped) > len(mapped_text):
                        remaining_part_raw = current_text_stripped[len(mapped_text):]
                        if remaining_part_raw.strip():
                            potential_split_info = (mapped_text, mapped_style, mapped_list_id)
                            break
                
                if potential_split_info:
                    mapped_text, mapped_style, mapped_list_id = potential_split_info
                    
                    split_offset = len(mapped_text)
                    
                    new_para = self._split_paragraph_at_offset(p, split_offset)
                    
                    if new_para is not None:
                        if mapped_style in doc.styles:
                            new_para.style = doc.styles[mapped_style]
                        else:
                            new_para.style = doc.styles["Normal"]
                            
                        if mapped_list_id is not None:
                            try:
                                new_para.list_id = mapped_list_id
                            except:
                                pass

                        p.style = doc.styles["Normal"]
                        if p._element.pPr is not None and not (mapped_style.startswith("Heading") or "Appendix" in mapped_style): 
                             numPr = p._element.pPr.find(qn('w:numPr'))
                             if numPr is not None:
                                 p._element.pPr.remove(numPr)
                        
                        # Strip leading whitespace from the original paragraph (Tail)
                        found_content = False
                        for run in p.runs:
                            if found_content:
                                break
                            text = run.text
                            if not text:
                                continue
                            if text.strip() == "":
                                run.text = ""
                            else:
                                run.text = text.lstrip()
                                found_content = True
                        
                        remaining_text = p.text.strip()
                        paragraphs_and_text[i] = (new_para, mapped_text)
                        paragraphs_and_text.insert(i + 1, (p, remaining_text))
                        
                        print(f"Split paragraph: '{mapped_text}' | '{remaining_text}'")
                        applied_style = True
                    else:
                        pass

                    
                elif not applied_style:
                    # Skip setting style if no match found - no fallback to Normal
                    pass
            
            i += 1

    def _split_paragraph_at_offset(self, paragraph, offset: int):
        """
        Split a paragraph at the specified character offset, preserving formatting.
        Returns the NEW paragraph (which contains content BEFORE the offset).
        The original paragraph retains content AFTER the offset.
        """
        if offset < 0:
            raise ValueError("Offset cannot be negative")
            
        full_text = paragraph.text
        if offset == 0:
            # Everything stays in original. Create empty new para before? 
            # Or effectively no split for 'before' content.
            # But the caller expects a new paragraph for the "Prefix".
            # If prefix is empty, maybe return None?
            # But for our specific use case (Heading split), offset will be > 0.
            return None
            
        if offset >= len(full_text):
            # Everything moves to new para. Original becomes empty?
            # We'll handle this standard logic.
            pass

        # We need to find exactly where to cut in the runs
        current_pos = 0
        split_run_index = -1
        split_in_run_offset = 0
        
        # 1. Identify which run contains the split point
        for idx, run in enumerate(paragraph.runs):
            run_len = len(run.text)
            if current_pos <= offset < current_pos + run_len:
                split_run_index = idx
                split_in_run_offset = offset - current_pos
                break
            current_pos += run_len
            
        if split_run_index == -1 and offset == len(full_text):
            # Split is at the very end
            split_run_index = len(paragraph.runs)
            split_in_run_offset = 0
        elif split_run_index == -1:
            # Should not happen if offset is valid
            return None

        # 2. Create the new paragraph BEFORE the current one
        new_paragraph = paragraph.insert_paragraph_before('')
        
        # 3. Move runs
        # We process runs in the original paragraph.
        # Runs BEFORE split_run_index: Move fully to new_paragraph
        # Run AT split_run_index: Split. First part to new, second stay (or re-add).
        # Runs AFTER split_run_index: Stay in original.
        
        # Careful: Modifying the list of runs while iterating or by index can be tricky 
        # because paragraph.runs is a dynamic property in python-docx.
        # Safer to collect what needs to be moved/changed, then apply.
        
        # Actually, standard python-docx `paragraph.runs` returns a list of proxy objects. 
        # Modifying the XML is the source of truth.
        
        # Strategy:
        # A. Copy content/formatting to new_paragraph for everything < offset
        # B. Remove that content from original paragraph
        
        runs_to_move = []
        split_run = None
        
        # We iterate a copy to be safe, though we aren't modifying *yet*
        all_runs = list(paragraph.runs)
        
        for i, run in enumerate(all_runs):
            if i < split_run_index:
                # Fully move
                new_run = new_paragraph.add_run(run.text)
                self._copy_run_formatting(run, new_run)
                # clear original
                run.text = ""
            elif i == split_run_index:
                # Split this run
                original_text = run.text
                text_before = original_text[:split_in_run_offset]
                text_after = original_text[split_in_run_offset:]
                
                if text_before:
                    new_run = new_paragraph.add_run(text_before)
                    self._copy_run_formatting(run, new_run)
                
                # Update original run to only have the 'after' text
                run.text = text_after
                
            else:
                # i > split_run_index
                # These stay in original. Do nothing.
                pass
                
        # Clean up empty runs in original? python-docx doesn't auto-delete empty runs usually, 
        # but visually it's fine. We can leave them or explicitly remove XML.
        # For robustness, let's leave them unless they cause issues.
        
        return new_paragraph
        
    def _copy_run_formatting(self, source_run, target_run):
        """Copy formatting from source_run to target_run."""
        if not source_run or not target_run:
            return
            
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        
        if not hasattr(source_run, 'font') or not hasattr(target_run, 'font'):
            return
            
        source_font = source_run.font
        target_font = target_run.font
        
        if hasattr(source_font, 'name'):
            target_font.name = source_font.name
        if hasattr(source_font, 'size'):
            target_font.size = source_font.size
            
        if hasattr(source_font, 'color') and source_font.color and source_font.color.rgb:
            if not hasattr(target_font, 'color'):
                target_font.color.rgb = source_font.color.rgb
            else:
                target_font.color.rgb = source_font.color.rgb
                
        font_properties = [
            'highlight_color', 'subscript', 'superscript', 'strike',
            'double_strike', 'shadow', 'outline', 'rtl', 'imprint',
            'cs_bold', 'complex_script', 'hidden'
        ]
        
        for prop in font_properties:
            if hasattr(source_font, prop):
                try:
                    setattr(target_font, prop, getattr(source_font, prop))
                except (AttributeError, KeyError):
                    continue

    def _format_references(self, doc, doc_structure):
        """
        Central method for handling all reference formatting operations.
        This includes:
        1. Using AI-detected references from doc_structure
        2. Applying reference-specific paragraph styles
        3. Sorting references alphabetically
        4. Applying text formatting patterns (italics, etc.)
        5. Setting hanging indentation and other reference-specific formatting
        """
        # 1. Extract AI-detected references from doc_structure (Linear Block Schema)
        ai_references = [
            block for block in doc_structure.get("blocks", []) 
            if block.get("type") == "reference_list_item"
        ]
        
        if not ai_references:
            return  # No references detected by AI
        
        # Normalize AI references to list of strings
        reference_texts = []
        for ref_entry in ai_references:
            content = ref_entry.get("content", "").strip()
            if content:
                reference_texts.append(content)
        
        if not reference_texts:
            return  # No valid reference texts
        
        # 2. Build a mapping of reference text to paragraph objects
        reference_paragraphs = []
        text_to_paragraph = {}
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                text_to_paragraph[text] = p
        
        # Match AI-detected references to actual paragraphs
        for ref_text in reference_texts:
            if ref_text in text_to_paragraph:
                reference_paragraphs.append(text_to_paragraph[ref_text])
        
        if not reference_paragraphs:
            return  # No matching paragraphs found
        
        # 3. Sort references alphabetically
        if len(reference_paragraphs) >= 2:
            def sort_key(p):
                """Sort key to handle numbered or unnumbered references."""
                text = p.text.strip()
                # Remove leading numbers, brackets, etc.
                cleaned_text = re.sub(r'^\s*[\(\[]?\d+[\)\]]?\.\?\s*', '', text)
                return cleaned_text.lower()

            sorted_paragraphs = sorted(reference_paragraphs, key=sort_key)
            
            # Only reorder if the order has changed
            if reference_paragraphs != sorted_paragraphs:
                # Store sorted texts
                sorted_texts = [p.text for p in sorted_paragraphs]
                # Update original paragraphs with sorted text
                for i, p in enumerate(reference_paragraphs):
                    p.text = sorted_texts[i]

        # 4. Apply reference formatting to each paragraph
        style_config = self.selected_style_guide["styles"].get("References", {})
        
        for paragraph in reference_paragraphs:
            # A. Apply the References style for basic paragraph formatting
            paragraph.style = doc.styles["References"]
            
            # B. Apply style-specific paragraph properties
            if "paragraph" in style_config:
                para_config = style_config["paragraph"]
                para_format = paragraph.paragraph_format
                
                # Set indentation (including hanging indent)
                if "left_indent" in para_config:
                    para_format.left_indent = para_config["left_indent"]
                if "first_line_indent" in para_config:
                    para_format.first_line_indent = para_config["first_line_indent"]
                
                # Set spacing
                if "space_before" in para_config:
                    para_format.space_before = para_config["space_before"]
                if "space_after" in para_config:
                    para_format.space_after = para_config["space_after"]
                if "line_spacing" in para_config:
                    para_format.line_spacing_rule = para_config["line_spacing"]

        # Removed regex-based inline formatting to prevent text corruption as requested.
        # Only paragraph-level styles and sorting are now applied to references.

    def _apply_font_properties(self, doc):
            """
            Explicitly applies required font properties (name, color, bold, size) to runs
            within paragraphs belonging to a specific set of target styles,
            overcoming python-docx's style application inconsistencies.
            """
            styles_config = self.selected_style_guide["styles"]
            
            # Define the set of styles that require this explicit fix (based on the original code)
            target_style_names = {
                'Normal', 'Abstract', 'Title', 'Appendix Title',
                'Appendices Title', 'List Bullet', 'List Number', 'List Item'
            }

            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                
                # 1. Check if the paragraph's style is one of the designated targets
                if style_name in target_style_names:
                    
                    # 2. Retrieve the font configuration for the applied style
                    style_config = styles_config.get(style_name)
                    
                    if style_config and "font" in style_config:
                        font_config = style_config["font"]
                        
                        # 3. Iterate over all runs in the paragraph to apply fixes
                        for run in paragraph.runs:
                            
                            # --- Apply Mandatory Font Properties (Name, Color, Bold) ---
                            # These are applied universally to all target styles
                            if "name" in font_config:
                                run.font.name = font_config["name"]
                            
                            if "color" in font_config:
                                run.font.color.rgb = font_config["color"]
                                
                            if "bold" in font_config:
                                run.font.bold = font_config["bold"]
                            
                            
                            # --- Apply Conditional Font Property (Size) ---
                            # Only apply size if the style is 'Normal' or 'Abstract', 
                            # mirroring the original logic's size constraint.
                            if style_name in {'Normal', 'Abstract'} and "size" in font_config:
                                run.font.size = font_config["size"]

    def _apply_explicit_paragraph_properties(self, doc):
        """
        Explicitly applies all essential paragraph formatting properties (spacing, 
        indentation, alignment, and keep options) from the style guide to each 
        paragraph's format object to ensure consistency and overcome python-docx limitations.
        """
        styles_config = self.selected_style_guide["styles"]
        
        # Check if "Appendices Title" (the master heading) exists in the document
        has_master_appendices = any(p.style.name == "Appendices Title" for p in doc.paragraphs)
        
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            # Normalize and handle known markdown-style paragraphs
            if isinstance(style_name, str) and style_name.lower() == 'ds-markdown-paragraph':
                # Convert to Normal style if available so paragraph-level properties are applied
                if 'Normal' in doc.styles:
                    try:
                        paragraph.style = doc.styles['Normal']
                        style_name = 'Normal'
                        print(f"Converted ds-markdown-paragraph to 'Normal' for paragraph: '{paragraph.text[:60]}...'")
                    except Exception:
                        # If conversion fails, fall back to original style_name
                        pass
            
            # Retrieve the configuration for the paragraph's style
            if style_name not in styles_config:
                continue
                
            style_config = styles_config[style_name]
            if not style_config or "paragraph" not in style_config:
                continue
                
            para_config = style_config["paragraph"]
            para_format = paragraph.paragraph_format
            
            # --- Apply Indentation and Alignment Properties ---
            if "alignment" in para_config:
                alignment = para_config["alignment"]
                
                # Dynamic alignment for Appendix Title: Left-align if master "Appendices Title" is present
                if style_name == "Appendix Title" and has_master_appendices:
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                para_format.alignment = alignment
            if "left_indent" in para_config:
                para_format.left_indent = para_config["left_indent"]
            if "right_indent" in para_config:
                para_format.right_indent = para_config["right_indent"]
            if "first_line_indent" in para_config:
                para_format.first_line_indent = para_config["first_line_indent"]
            
            if "space_before" in para_config:
                para_format.space_before = para_config["space_before"]
            if "space_after" in para_config:
                para_format.space_after = para_config["space_after"]
            if "line_spacing" in para_config:
                para_format.line_spacing = para_config["line_spacing"]

            # --- Apply Flow Control Properties ---
            if "keep_together" in para_config:
                para_format.keep_together = para_config["keep_together"]
            if "keep_with_next" in para_config:
                para_format.keep_with_next = para_config["keep_with_next"]
            if "page_break_before" in para_config:
                para_format.page_break_before = para_config["page_break_before"]
                
            # --- Apply Widow/Orphan Control ---
            if "widow_control" in para_config:
                para_format.widow_control = para_config["widow_control"]
            if "orphan_control" in para_config:
                try:
                    para_format.orphan_control = para_config["orphan_control"]
                except AttributeError:
                    # orphan_control may not be available in all python-docx versions
                    pass
            
            # --- Apply Outline Level ---
            if "outline_level" in para_config:
                try:
                    outline_level = para_config["outline_level"]
                    
                    # Demote outline level if master "Appendices Title" is present (Appendix Title becomes Level 1)
                    if style_name == "Appendix Title" and has_master_appendices:
                        outline_level = 1
                        
                    para_format.outline_level = outline_level
                except AttributeError:
                    # outline_level may not be available in all python-docx versions
                    pass

    def _create_abstract_num(self, numbering_part, numbering_type):
        """Creates a new abstractNum element in the numbering part."""
        abstract_nums = numbering_part._element.xpath('//w:abstractNum')
        new_id = 0
        if abstract_nums:
            ids = [int(n.get(qn('w:abstractNumId'))) for n in abstract_nums]
            new_id = max(ids) + 1

        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), str(new_id))

        # Add MultiLevelType
        multiLevelType = OxmlElement('w:multiLevelType')
        multiLevelType.set(qn('w:val'), 'hybridMultilevel')
        abstractNum.append(multiLevelType)

        # Level definition
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), numbering_type)
        lvl.append(numFmt)

        lvlText = OxmlElement('w:lvlText')
        if numbering_type == 'bullet':
            lvlText.set(qn('w:val'), '·') # Standard bullet
            lvl.append(lvlText)
            
            # For bullets, we NEED to set the font to Symbol or wingdings
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Symbol')
            rFonts.set(qn('w:hAnsi'), 'Symbol')
            rFonts.set(qn('w:hint'), 'default')
            rPr.append(rFonts)
            lvl.append(rPr)
        elif numbering_type == 'lowerLetter':
            lvlText.set(qn('w:val'), '%1.')
            lvl.append(lvlText)
        else: # decimal
            lvlText.set(qn('w:val'), '%1.')
            lvl.append(lvlText)

        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)

        abstractNum.append(lvl)

        # Insert at the beginning or after other abstractNums
        if abstract_nums:
            abstract_nums[-1].addnext(abstractNum)
        else:
            numbering_part._element.insert(0, abstractNum)

        return new_id

    def _apply_list_properties(self, doc):
        """
        Enhanced version of list property application.
        Uses dynamic abstract numbering to ensure restarts work correctly.
        """
        numbering_part = self._ensure_numbering_part(doc)
        if not numbering_part:
            return

        # Cache for num_ids associated with our (type, list_id) pairs
        # Keys: (numbering_type, ai_list_id)
        list_id_map = {}

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name
            
            # Check if this paragraph should be treated as a list item
            # We check both the style name AND if it has a list_id assigned (from _format_content_in_place)
            ai_list_id = getattr(paragraph, 'list_id', None)
            is_list_style = style_name in {"List Bullet", "List Number", "List Alphabet", "List Item"}
            
            if not is_list_style and ai_list_id is None:
                continue

            # Determine numbering type
            if style_name == "List Bullet":
                num_type = 'bullet'
            elif style_name == "List Alphabet":
                num_type = 'lowerLetter'
            else:
                num_type = 'decimal'

            # If we have an ai_list_id, use it to group and potentially restart
            # If not, treat as a single continuous list for that type (legacy behavior)
            group_key = (num_type, ai_list_id if ai_list_id is not None else "default")

            if group_key not in list_id_map:
                # Need to create a new num mapping
                abs_id = self._create_abstract_num(numbering_part, num_type)
                
                # Create the <w:num> element that points to our abstractNum
                num_nums = numbering_part._element.xpath('//w:num')
                new_num_id = 1
                if num_nums:
                    num_ids = [int(n.get(qn('w:numId'))) for n in num_nums]
                    new_num_id = max(num_ids) + 1
                
                num_el = OxmlElement('w:num')
                num_el.set(qn('w:numId'), str(new_num_id))
                abs_ref = OxmlElement('w:abstractNumId')
                abs_ref.set(qn('w:val'), str(abs_id))
                num_el.append(abs_ref)
                
                if num_nums:
                    num_nums[-1].addnext(num_el)
                else:
                    numbering_part._element.append(num_el)
                
                list_id_map[group_key] = new_num_id

            target_num_id = list_id_map[group_key]

            # Apply to paragraph XML
            try:
                p_el = paragraph._p
                pPr = p_el.get_or_add_pPr()
                numPr = pPr.get_or_add_numPr()
                
                ilvl = numPr.get_or_add_ilvl()
                ilvl.set(qn('w:val'), '0')
                
                numId = numPr.get_or_add_numId()
                numId.set(qn('w:val'), str(target_num_id))
            except Exception as e:
                print(f"Warning: Failed to set list properties for paragraph: {e}")

            # Ensure runs are black (explicit requirement)
            for run in paragraph.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    continue

    def _ensure_numbering_part(self, doc):
        """
        Ensure the document has a numbering part and return the numbering element.
        Create a minimal numbering.xml structure if it doesn't exist.
        """
        try:
            numbering = doc.part.numbering_part
            return numbering
        except Exception:
            # python-docx should create numbering_part when needed, but guard anyway
            return None
    
    def _apply_heading_styles(self, doc):
            """Apply heading styles based on the style guide."""
            for paragraph in doc.paragraphs:
                
                if paragraph.style.name.startswith('Heading') or paragraph.style.name in {'Appendix Title', 'Appendices Title'}:
                    style_name = paragraph.style.name
                    if style_name.startswith('Heading'):
                        heading_level = int(style_name.split(' ')[1])
                        target_style_key = f"Heading {heading_level}"
                    else:
                        target_style_key = style_name

                    # *** REDUNDANCY FIXED HERE: Cache the full font configuration ***
                    font_config = STYLE_GUIDES[self.selected_style_name]['styles'][target_style_key]['font']

                    for run in paragraph.runs:
                        # Now use the cached variable, eliminating repeated lookups
                        run.font.name = font_config['name']
                        run.font.size = font_config['size']
                        run.font.bold = font_config.get('bold')
                        run.font.italic = font_config.get('italic')
                        run.font.color.rgb = font_config.get('color')
                        run.font.underline = font_config.get('underline')
                        run.font.space_before = font_config.get('space_before')
                        run.font.space_after = font_config.get('space_after')

    def _remove_blank_lines(self, doc, boundary_idx=0):
        """Remove empty paragraphs after the boundary index."""
        to_remove = [p for p in doc.paragraphs[boundary_idx:] 
                     if not p.text.strip() and not p.runs]
        
        for p in reversed(to_remove):
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

    def _add_page_numbers(self, doc, has_title_page, title_page_boundary):
        from copy import deepcopy
        
        position = self.selected_style_guide["meta"].get("page_numbers")
        if not position:
            return

        if has_title_page:
            # 1. Ensure document has at least two sections for title page and body.
            if len(doc.sections) == 1 and title_page_boundary > 0:
                # Add a section break by inserting section properties into the last paragraph of the title page.
                last_para_of_title = doc.paragraphs[title_page_boundary - 1]
                pPr = last_para_of_title._p.get_or_add_pPr()
                pPr.append(deepcopy(doc.sections[0]._sectPr))

            # 2. Configure sections.
            # Section 1 (Title Page): Unnumbered.
            title_section = doc.sections[0]
            title_section.different_first_page_header_footer = True
            # Clear headers/footers for the title page itself.
            for container in [title_section.first_page_header, title_section.first_page_footer]:
                for p in list(container.paragraphs):
                    p._element.getparent().remove(p._element)

            # Section 2 (Body): Starts numbering at 1.
            if len(doc.sections) > 1:
                body_section = doc.sections[1]
                # Set page number to start at 1 and use Arabic numerals.
                pgNumType = body_section._sectPr.get_or_add_pgNumType()
                pgNumType.set(qn('w:start'), '1')
                pgNumType.set(qn('w:fmt'), 'decimal')
                section_for_numbering = body_section
            else:
                # Fallback if section break failed; use primary header of the first section.
                section_for_numbering = doc.sections[0]
        else:
            # Document has no title page; use the first and only section.
            section_for_numbering = doc.sections[0]

        # 3. Add page number field to the correct container.
        if "header" in position:
            container = section_for_numbering.header
        elif "footer" in position:
            container = section_for_numbering.footer
        else:
            return  # Unknown position.

        # Clear any existing content from the target container.
        for p in list(container.paragraphs):
            p._element.getparent().remove(p._element)

        paragraph = container.add_paragraph()
        
        if "right" in position:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "center" in position:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 4. Insert the PAGE field using OXML elements.
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar_begin, instrText, fldChar_end])




def validate_style(style_name: str) -> str:
    """Validate and return the style name, defaulting to APA if invalid."""
    return style_name if style_name.lower() in STYLE_GUIDES else "apa"

from pathlib import Path

def determine_output_path(input_path: Path, output_arg: str, style_name: str) -> Path:
    """Determine the output path based on input and arguments."""
    if output_arg:
        return Path(output_arg)
    
    formatting_dir = Path("formatted")
    formatting_dir.mkdir(exist_ok=True)

    output_path = formatting_dir / f"{input_path.stem}_formatted_{style_name}{input_path.suffix}"
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")

    return output_path


def track_changes(input_path: Path, output_path: Path):
    """Track changes in document."""
    from utils.track_changes import TrackChanges
    import os
    
    # Convert to absolute paths
    input_file = str(input_path.absolute())
    output_file = str(output_path.absolute())
    
    # Debug output
    print(f"Original file: {input_file}")
    print(f"Formatted file: {output_file}")
    
    # Initialize and run the tracker
    tracker = TrackChanges(input_file, output_file)
    tracker.compare_docs()

def list_available_styles():
    """List all available citation styles."""
    print("Available citation styles:")
    for style_name, style_config in STYLE_GUIDES.items():
        meta = style_config.get("meta", {})
        print(f"  {style_name.upper():8} - {meta.get('default_font', 'Times New Roman')} font")
        print(f"          {'Title page: ' + ('Yes' if meta.get('title_page') else 'No'):20}")
        print(f"          {'Abstract: ' + ('Required' if meta.get('requires_abstract') else 'Optional'):20}")
        print()

def get_input_file_interactive():
    """Prompt user for input file in interactive mode."""
    print("\nNote: If you provide just a filename, it will be searched in the 'documents' folder.")
    while True:
        input_file = input("Enter the path to your Word document: ").strip()
        if not input_file:
            print("Please provide a file path.")
            continue
        
        input_path = Path(input_file)
        
        if input_path.exists():
            if input_path.suffix.lower() == '.docx':
                return input_path
            else:
                print("Error: File must be a Word document (.docx)")
        else:
            documents_path = Path("documents") / input_path.name
            if documents_path.exists():
                if documents_path.suffix.lower() == '.docx':
                    return documents_path
                else:
                    print("Error: File must be a Word document (.docx)")
            else:
                print(f"Error: File '{input_file}' not found (also checked in 'documents' folder).")

def validate_input_file(input_path: str) -> Path:
    """Validate input file exists and is a Word document."""
    if not input_path:
        raise ValueError("Input file path is required")
    
    path = Path(input_path)
    
    if not path.exists():
        documents_path = Path("documents") / path.name
        if documents_path.exists():
            path = documents_path
        else:
            raise FileNotFoundError(f"Input file not found: {input_path} (also checked in 'documents' folder)")
    
    if path.suffix.lower() != '.docx':
        raise ValueError(f"Input file must be a Word document (.docx), got: {path.suffix}")
    
    return path

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Format documents according to citation styles.')
    
    parser.add_argument('input', nargs='?', help='Input Word document path')
    
    parser.add_argument('-s', '--style', default='apa',
                      help='Citation style (default: apa). Use --list-styles to see available options.')
    parser.add_argument('-o', '--output', help='Output file path (default: input_filename_formatted_style.docx)')
    parser.add_argument('--english', choices=['us', 'gb', 'au', 'ca'], default='us',
                      help='English language variant for spelling (default: us)')
    
    parser.add_argument('-i', '--interactive', action='store_true',
                      help='Interactive mode - prompts for input file if not provided')
    parser.add_argument('-l', '--list-styles', action='store_true',
                      help='List available citation styles and exit')
    parser.add_argument('-t', '--track-changes', action='store_true',
                        help='Enable track changes in document')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--report-only', action='store_true', 
                      help='Generate comprehensive reports (spelling, formatting) without modifying the document.')
    group.add_argument('--fix-errors', action='store_true', 
                      help='Automatically fix spelling, punctuation, and capitalization errors, then format the document.')
    group.add_argument('--format', action='store_true', 
                      help='Apply document formatting based on the chosen style, skipping error correction.')
    parser.add_argument('--overwrite', action='store_true', help='Overwrite output file if it exists')

    mla_group = parser.add_mutually_exclusive_group()
    mla_group.add_argument('--mla-title-page', action='store_true',
                          help='For MLA style, generate a title page (no page number on the title page).')
    mla_group.add_argument('--mla-heading', action='store_true',
                          help='For MLA style, use a first-page heading instead of a title page (page numbering starts from page 1).')

    return parser.parse_args()

def main():
    args = parse_arguments()
    
    if args.list_styles:
        list_available_styles()
        return 0
    
    input_path = None
    if args.input:
        try:
            input_path = validate_input_file(args.input)
        except (FileNotFoundError, ValueError) as e:
            print(f"Error: {e}")
            return 1
    elif args.interactive:
        try:
            input_path = get_input_file_interactive()
        except KeyboardInterrupt:
            print("\nOperation cancelled by user.")
            return 1
    else:
        print("Error: Input file is required. Use --interactive mode or provide a file path.")
        print("Usage: python app.py input.docx [options]")
        print("   or: python app.py --interactive [options]")
        return 1
    
    style_name = validate_style(args.style)
    
    try:
        doc = Document(str(input_path))
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
        
        if args.report_only:
            print(f"Generating comprehensive report for: {input_path}")
            checker = DocumentChecker(english_variation=args.english)
            formatter_analyzer = FormattingAnalyzer(style_name)
            try:
                report = checker.get_correction_report(paragraphs_text)
                print("\n=== SPELLING REPORT ===")
                print(format_error_report(report))
                
                print("\n=== FORMATTING ANALYSIS ===")
                analysis_result = formatter_analyzer.analyze_document(str(input_path))
                analysis_report = formatter_analyzer.generate_report(analysis_result)
                print(analysis_report)
                
                print("\nReport generated successfully. No changes made to document.")
                return 0
            except Exception as e:
                print(f"Error during report generation: {e}")
                return 1
            finally:
                checker.close()
        
        elif args.fix_errors:
            print(f"Auto-fixing spelling, punctuation, and capitalization errors: {input_path}")
            auto_corrector = AutoCorrector(english_variation=args.english)
            gemini_corrector = None
            try:
                print("\nInitializing Gemini AI for comprehensive text correction...")
                gemini_corrector = GeminiCorrector()
                print("Gemini AI initialized successfully.")
            except Exception as e:
                print(f"Warning: Could not initialize Gemini corrector: {e}")
                print("Falling back to basic spell checker.")

            try:
                if gemini_corrector:
                    print("\nUsing Gemini AI to apply comprehensive text corrections...")
                    print("   Simultaneous correction of spelling, punctuation, and capitalization in progress...")
                    corrected_paragraphs, correction_summary = gemini_corrector.correct_paragraphs(paragraphs_text)
                else:
                    print("\nApplying basic automatic corrections...")
                    print("   Note: Only basic spelling will be checked without Gemini AI")
                    corrected_paragraphs, correction_summary = auto_corrector.apply_all(paragraphs_text)
                
                idx = 0
                for p in doc.paragraphs:
                    if p.text.strip() and idx < len(corrected_paragraphs):
                        p.text = corrected_paragraphs[idx]
                        idx += 1
                
                # Display correction summary
                if correction_summary['total_corrections'] > 0:
                    print(auto_corrector.generate_correction_explanation(correction_summary))
                    # Get breakdown by type
                    by_type = correction_summary.get("corrections_by_type", {})
                    print(f"✅ Successfully applied {correction_summary['total_corrections']} corrections:")
                    print(f"   🔤 Spelling: {by_type.get('spelling', 0)}")
                    print(f"   🔣 Punctuation: {by_type.get('punctuation', 0)}")
                    print(f"   🔠 Capitalization: {by_type.get('capitalization', 0)}")
                else:
                    print("✅ No issues found that needed correction.")
                    
                print("\n📝 Proceeding with document formatting...")
            except Exception as e:
                print(f"Error during auto-fix: {e}")
                return 1
            finally:
                if auto_corrector:
                    auto_corrector.close()
                if gemini_corrector:
                    del gemini_corrector
        
        elif args.format:
            print(f"🎨 Applying formatting only (skipping error correction): {input_path}")
            # No spelling/grammar checks or corrections are performed in this mode.
            # The document object 'doc' remains as loaded from input_path.
            
        # Only proceed with formatting if not in report-only mode
        if not args.report_only:
            # Check API key for formatting
            # if not API_KEY:
            #     print("Error: GEMINI_API_KEY environment variable not set.")
            #     print("Please add your Gemini API key to the .env file.")
            #     return 1
            
            print(f"📝 Formatting document: {input_path}")
            print(f"🎨 Style: {style_name.upper()}, English: {args.english.capitalize()}")
            
            formatter = AdvancedFormatter(style_name)
            # Propagate MLA-specific options to the formatter
            if style_name == 'mla':
                setattr(formatter, 'mla_heading', bool(getattr(args, 'mla_heading', False)))
                setattr(formatter, 'mla_title_page', bool(getattr(args, 'mla_title_page', False)))
            output_path = determine_output_path(input_path, args.output, style_name)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Check if output file already exists
            if output_path.exists() and not args.overwrite:
                response = input(f"Output file '{output_path}' already exists. Overwrite? (y/N): ")
                if response.lower() not in ['y', 'yes']:
                    print("Operation cancelled.")
                    return 0
            
            # If --fix-errors was used, the 'doc' object already contains the corrected text.
            formatter.format_document(str(input_path), str(output_path))  # Pass original path, formatter will load it
            if args.track_changes:
                print(f"✅ Document formatted with Tracked Changes successfully: {output_path}")
                # Play a sound to indicate completion
                winsound.PlaySound("Beep", winsound.SND_ALIAS)
                track_changes(input_path, output_path)
            else:
                print(f"✅ Document formatted successfully: {output_path}")
                # Play a sound to indicate completion
                winsound.PlaySound("Beep", winsound.SND_ALIAS)
                os.startfile(output_path)
                
    except Exception as e:
        print(f"Error: Failed to format document: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for debugging
        print("Please check your input file and try again.")
        return 1

if __name__ == "__main__":
    raise SystemExit(main())



